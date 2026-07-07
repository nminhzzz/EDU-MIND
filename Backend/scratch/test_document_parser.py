import requests
import io
from docx import Document

BASE_URL = "http://127.0.0.1:8000"


def test_document_parser():
    print("--- 1. LOGGING IN AS TEACHER ---")
    login_res = requests.post(
        f"{BASE_URL}/api/v1/auth/login",
        json={"email": "teacher@edumind.com", "password": "teacherpassword"},
    )
    token = login_res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    print("\n--- 2. CREATING A MOCK DOCX WITH RICH CONTENT ---")
    # Sử dụng python-docx tạo tài liệu thực tế chứa nội dung kiến thức
    doc = Document()
    doc.add_heading("Triết học Mác - Lênin về Vật chất và Ý thức", level=1)
    doc.add_paragraph(
        "Vật chất là một phạm trù triết học dùng để chỉ thực tại khách quan được đem lại cho con người trong cảm giác."
    )
    doc.add_paragraph(
        "Ý thức là sự phản ánh năng động, sáng tạo thế giới khách quan vào bộ óc con người."
    )

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()

    files = {
        "file": (
            "test_tutor_material.docx",
            file_stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    data = {"subject_id": 1, "title": "Chuyên đề Vật chất và Ý thức"}

    print("\n--- 3. UPLOADING DOCX TO BACKEND ---")
    upload_res = requests.post(
        f"{BASE_URL}/api/v1/documents/", data=data, files=files, headers=headers
    )

    print(f"Upload Status Code: {upload_res.status_code}")
    assert upload_res.status_code == 201
    doc_id = upload_res.json()["id"]
    print(f"Document created successfully with ID: {doc_id}")

    print("\n--- 4. VERIFYING EXTRACTED CONTENT IN MONGODB ---")
    # Kết nối MongoDB trực tiếp qua API hoặc query mongo để xem RAG content
    from app.database.mongodb import get_mongodb_db
    import asyncio

    async def verify_mongo():
        db_mongo = get_mongodb_db()
        doc_mongo = await db_mongo.study_material_embeddings.find_one(
            {"metadata.document_id": doc_id}
        )
        assert doc_mongo is not None
        content = doc_mongo["content"]
        print(f"Extracted content found in MongoDB:\n{content}\n")
        assert "Vật chất là một phạm trù" in content
        assert "Ý thức là sự phản ánh" in content
        print("🎉 SUCCESS: Document content was parsed and stored accurately for RAG!")

    asyncio.run(verify_mongo())

    print("\n--- 5. CLEANING UP DOCUMENT ---")
    delete_res = requests.delete(
        f"{BASE_URL}/api/v1/documents/{doc_id}", headers=headers
    )
    print(f"Delete Status Code: {delete_res.status_code}")
    assert delete_res.status_code == 200
    print("Cleanup successful!")


if __name__ == "__main__":
    test_document_parser()
