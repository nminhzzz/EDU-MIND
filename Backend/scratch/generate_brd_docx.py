"""
Script tạo file BRD (.docx) cho dự án EDU-MIND
Bỏ toàn bộ code blocks, giữ nội dung nghiệp vụ + bảng + định dạng Word chuẩn.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Cấu hình page ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 51, 102)
    if i == 1:
        hs.font.size = Pt(18)
    elif i == 2:
        hs.font.size = Pt(16)
    elif i == 3:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(13)


def add_table(headers, rows, col_widths=None):
    """Tạo bảng Word có viền + header đậm."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_para(text, bold=False, italic=False, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def add_note(text, label="Lưu ý"):
    p = doc.add_paragraph()
    run_label = p.add_run(f"⚠️ {label}: ")
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = 'Times New Roman'
    run_label.font.color.rgb = RGBColor(180, 0, 0)
    run_text = p.add_run(text)
    run_text.font.size = Pt(12)
    run_text.font.name = 'Times New Roman'
    run_text.italic = True


# ═══════════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BUSINESS REQUIREMENTS DOCUMENT (BRD)")
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("HỆ THỐNG HỌC TẬP THÔNG MINH — EDU-MIND")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
doc.add_paragraph()

add_table(
    ["Thông tin", "Chi tiết"],
    [
        ["Tên dự án", "EDU-MIND — Smart Learning Assistant Platform"],
        ["Phiên bản", "2.0.0"],
        ["Ngày tạo", "2026-06-25"],
        ["Tác giả", "[Product Owner]"],
        ["Trạng thái", "Final Draft"],
    ],
    col_widths=[6, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# MỤC LỤC (placeholder)
# ═══════════════════════════════════════════════════════════════
doc.add_heading("MỤC LỤC", level=1)
toc_items = [
    "1. Project Overview",
    "2. Problem Statement",
    "3. Business Objectives",
    "4. Stakeholders",
    "5. Scope",
    "6. Functional Requirements",
    "   6.1. Phân hệ Giáo viên (Teacher Web Portal)",
    "   6.2. Phân hệ Học sinh (Student Mobile/Web App)",
    "   6.3. Phân hệ Quản trị viên (System Admin Portal)",
    "7. Non-Functional Requirements",
    "8. Success Metrics",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading("1. PROJECT OVERVIEW", level=1)

doc.add_heading("1.1. Tổng quan dự án", level=2)
add_para(
    "EDU-MIND là một nền tảng học tập thông minh (Smart Learning Assistant Platform) kết hợp sức mạnh của "
    "AI Agents với hệ thống quản lý lớp học truyền thống, nhằm cá nhân hóa toàn bộ trải nghiệm tự học của "
    "học sinh — từ lộ trình học tập, gia sư ảo hỏi đáp, đến tự luyện đề thi và đánh giá năng lực thích ứng."
)
add_para(
    "Hệ thống phục vụ 3 nhóm người dùng chính: Học sinh (Student), Giáo viên (Teacher), và Quản trị viên (Admin), "
    "vận hành trên kiến trúc cơ sở dữ liệu Hybrid (MySQL + MongoDB) cùng lớp AI Agents tự động hóa."
)

doc.add_heading("1.2. Kiến trúc kỹ thuật tổng quan", level=2)
add_table(
    ["Tầng", "Công nghệ"],
    [
        ["Frontend", "Next.js (React), TailwindCSS, Axios, SSR/CSR"],
        ["Backend", "FastAPI (Python 3.12), Pydantic, SQLAlchemy"],
        ["Database — SQL", "MySQL — 14 bảng quan hệ (toàn vẹn ACID)"],
        ["Database — NoSQL", "MongoDB — 3 collections chính + 1 log (Vector RAG)"],
        ["Cache", "Redis — JWT token cache, draft lộ trình"],
        ["AI / LLM", "OpenAI API, LangChain, Mastra Agent — 5 AI Agents"],
        ["Realtime", "SSE (Server-Sent Events) — streaming chat AI"],
        ["File Storage", "Cloudinary (cloud) / Local fallback"],
        ["DevOps", "Docker, Docker Compose, Nginx reverse proxy"],
    ],
    col_widths=[5, 13]
)

doc.add_heading("1.3. Tổng quan cơ sở dữ liệu", level=2)
add_para("Hệ thống sử dụng kiến trúc cơ sở dữ liệu Hybrid gồm:", bold=True)
add_bullet("MySQL (14 bảng): Lưu trữ các thực thể có cấu trúc chặt chẽ, ràng buộc toàn vẹn dữ liệu cao.")
add_bullet("MongoDB (3 collections + 1 log): Lưu dữ liệu phi cấu trúc, chat logs, Vector Embeddings RAG.")
add_bullet("Redis: Cache JWT token và draft lộ trình học tập.")

add_para("Danh sách 14 bảng MySQL:", bold=True)
add_table(
    ["#", "Tên bảng", "Mục đích"],
    [
        ["1", "users", "Thông tin tài khoản (Học sinh, Giáo viên, Admin) + Enum grade khối lớp"],
        ["2", "student_preferences", "Cấu hình thời gian rảnh & cam kết học tập (1-1 với users)"],
        ["3", "classrooms", "Lớp học do Giáo viên phụ trách, gắn với 1 môn học"],
        ["4", "classroom_students", "Bảng N-N liên kết Học sinh ↔ Lớp học"],
        ["5", "subjects", "Danh mục môn học chính thức trên hệ thống"],
        ["6", "study_documents", "Metadata tài liệu giảng dạy (PDF/DOCX/TXT/MD)"],
        ["7", "study_goals", "Mục tiêu điểm số + deadline cho từng môn"],
        ["8", "study_plans", "Nhiệm vụ học tập hàng ngày do AI phân bổ"],
        ["9", "study_plan_progress", "Tiến độ % hoàn thành từng nhiệm vụ ngày"],
        ["10", "quizzes", "Đề thi trắc nghiệm JSON do AI sinh từ tài liệu RAG"],
        ["11", "quiz_attempts", "Kết quả chấm bài tự động + lịch sử làm bài"],
        ["12", "learning_analytics", "Hồ sơ học lực AI: điểm mạnh/yếu theo chủ đề"],
        ["13", "ai_recommendation_reviews", "Đề xuất ôn tập AI đã tự duyệt (HITL Bypass)"],
        ["14", "notifications", "Thông báo đẩy gửi tới người dùng"],
    ],
    col_widths=[1, 5.5, 11.5]
)

add_para("Danh sách MongoDB Collections:", bold=True)
add_table(
    ["#", "Tên collection", "Mục đích"],
    [
        ["1", "chat_sessions", "Phiên hội thoại giữa học sinh và AI Gia sư ảo"],
        ["2", "chat_messages", "Chi tiết tin nhắn trong phiên chat (bộ nhớ hội thoại)"],
        ["3", "study_material_embeddings", "Vector Store RAG — chunks tài liệu + embedding 4096 chiều"],
        ["4", "ai_logs", "Log cuộc gọi LLM: input, output, execution_time, agent_name"],
    ],
    col_widths=[1, 6.5, 10.5]
)

add_para(
    "Nguyên tắc Hybrid Database: MySQL chịu trách nhiệm toàn vẹn dữ liệu (ACID, FK constraints). "
    "MongoDB chịu trách nhiệm dữ liệu phi cấu trúc. Liên kết cross-database qua student_id (Int) "
    "và subject_id (Int) — tham chiếu logic, KHÔNG phải Foreign Key vật lý.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading("2. PROBLEM STATEMENT", level=1)

doc.add_heading("2.1. Vấn đề của Học sinh", level=2)
add_table(
    ["#", "Vấn đề", "Hệ quả"],
    [
        ["P1", "Không có lộ trình học tập cá nhân hóa", "Học tràn lan, mất phương hướng trước deadline"],
        ["P2", "Không có ai hỏi đáp ngoài giờ học", "Gặp bài khó lúc 22h tối → bỏ qua kiến thức"],
        ["P3", "Không biết điểm yếu ở đâu", "Chỉ biết điểm tổng, không biết hổng chương nào"],
        ["P4", "Tự ôn tập thiếu kỷ luật", "Không có nhắc nhở → học dồn trước thi"],
    ],
    col_widths=[1, 6, 11]
)

doc.add_heading("2.2. Vấn đề của Giáo viên", level=2)
add_table(
    ["#", "Vấn đề", "Hệ quả"],
    [
        ["P5", "Soạn đề thi tốn rất nhiều thời gian", "Trung bình 2–4 giờ cho 1 bộ đề 30 câu có giải thích"],
        ["P6", "Chấm bài thủ công, không có phân tích", "Chỉ biết điểm tổng, không biết lớp yếu chủ đề nào"],
        ["P7", "Không thể cá nhân hóa bài tập cho từng HS", "1 GV quản lý 100+ HS, không thể giao bài riêng"],
    ],
    col_widths=[1, 7, 10]
)

doc.add_heading("2.3. Vấn đề của Hệ thống truyền thống", level=2)
add_table(
    ["#", "Vấn đề", "Hệ quả"],
    [
        ["P8", "Tài liệu phân tán, không có kho tri thức tập trung", "HS tự tìm tài liệu ngoài → chất lượng không đảm bảo"],
        ["P9", "Duyệt thủ công (HITL) làm chậm quá trình ôn tập", "GV phải duyệt từng đề xuất AI → HS chờ đợi"],
    ],
    col_widths=[1, 8, 9]
)

doc.add_heading("2.4. Phân tích khoảng cách (Gap Analysis)", level=2)
add_table(
    ["Hiện trạng (AS-IS)", "Khoảng cách (GAP)", "Mục tiêu (TO-BE) — EDU-MIND"],
    [
        ["HS tự tìm tài liệu", "Thiếu hệ thống RAG", "AI trả lời từ kho tài liệu GV (Vector Search)"],
        ["GV soạn đề thủ công", "Thiếu Quiz Generator AI", "AI tự sinh đề thi từ tài liệu RAG"],
        ["Chấm bài → chỉ có điểm số", "Thiếu Analytics AI", "AI phân tích weak/strong topics từng HS"],
        ["Muốn ôn tập → chờ GV duyệt", "Thiếu Auto-Approve", "AI tự duyệt + tạo lịch ôn tập (HITL Bypass)"],
    ],
    col_widths=[5, 4.5, 8.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. BUSINESS OBJECTIVES
# ═══════════════════════════════════════════════════════════════
doc.add_heading("3. BUSINESS OBJECTIVES", level=1)

doc.add_heading("3.1. Mục tiêu đo lường cụ thể (OKR)", level=2)
add_table(
    ["#", "Mục tiêu kinh doanh", "Chỉ số đo lường (KPI)", "Mục tiêu"],
    [
        ["OBJ-01", "Tăng hiệu quả tự học của Học sinh", "Tỷ lệ hoàn thành lộ trình (study_plans done / tổng)", "≥ 70%"],
        ["OBJ-02", "Giảm thời gian soạn đề cho GV", "Thời gian từ bấm 'Tạo đề' đến khi có đề hoàn chỉnh", "≤ 30 giây"],
        ["OBJ-03", "Cá nhân hóa ôn tập dựa trên điểm yếu", "% HS nhận lịch ôn tập AI bổ sung khi điểm < 8.0", "100% (tự động)"],
        ["OBJ-04", "Loại bỏ nghẽn cổ chai duyệt thủ công", "Thời gian từ AI đề xuất → HS nhận lịch ôn tập", "≤ 5 giây"],
        ["OBJ-05", "Cung cấp AI Gia sư 24/7", "Tỷ lệ câu hỏi AI trả lời chính xác (dựa trên RAG)", "≥ 85%"],
        ["OBJ-06", "Giúp GV nhận diện lỗ hổng tập thể", "% lớp có báo cáo common_weak_topics tự động", "100%"],
    ],
    col_widths=[2, 5.5, 6, 4.5]
)

doc.add_heading("3.2. Giá trị kinh doanh theo nhóm người dùng", level=2)
add_para("Đối với Học sinh:", bold=True)
add_bullet("Lộ trình cá nhân hóa theo lịch rảnh và cam kết học tập cá nhân.")
add_bullet("Gia sư AI hỏi đáp 24/7 bám sát tài liệu giảng dạy của giáo viên (RAG).")
add_bullet("Tự luyện đề thi trắc nghiệm + chấm điểm tức thì + giải thích từng câu.")
add_bullet("Nhận lịch ôn tập bổ sung TỰ ĐỘNG khi yếu — không cần chờ giáo viên duyệt.")

add_para("Đối với Giáo viên:", bold=True)
add_bullet("Giảm 70% thời gian soạn đề thi trắc nghiệm.")
add_bullet("Dashboard phân tích học lực cá nhân và tập thể — AI tự động cập nhật.")
add_bullet("Báo cáo lỗ hổng kiến thức tập thể cả lớp (common_weak_topics).")
add_bullet("Không cần duyệt thủ công đề xuất AI (HITL Bypass).")

add_para("Đối với Quản trị viên:", bold=True)
add_bullet("Giám sát chi phí token AI theo từng Agent.")
add_bullet("Quản lý người dùng & danh mục môn học tập trung.")
add_bullet("Dọn dẹp log AI cũ tự động để tiết kiệm dung lượng MongoDB.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. STAKEHOLDERS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("4. STAKEHOLDERS", level=1)

doc.add_heading("4.1. Ma trận vai trò & quyền hạn", level=2)
add_table(
    ["Vai trò", "Mô tả", "Phân hệ truy cập", "Quyền dữ liệu"],
    [
        ["Học sinh (Student)", "Người tham gia học tập, làm bài thi tự đánh giá", "Student App", "Chỉ đọc/ghi dữ liệu CỦA MÌNH"],
        ["Giáo viên (Teacher)", "Quản lý lớp học, tải tài liệu, theo dõi học lực", "Teacher Portal", "Đọc/ghi lớp CỦA MÌNH, đọc analytics HS trong lớp"],
        ["Quản trị viên (Admin)", "Vận hành hệ thống", "Admin Portal", "Đọc/ghi TOÀN BỘ hệ thống"],
        ["AI Agent (Hệ thống)", "Tác tử AI tự động hóa nghiệp vụ", "Background Tasks", "Đọc MySQL + MongoDB, ghi study_plans, quizzes, notifications..."],
    ],
    col_widths=[4, 5, 3.5, 5.5]
)

doc.add_heading("4.2. Chi tiết 5 AI Agents", level=2)
add_table(
    ["#", "Tên Agent", "Vai trò nghiệp vụ", "Dữ liệu tương tác"],
    [
        ["1", "Goal Planner", "Lập lộ trình học tập hàng ngày dựa trên mục tiêu + lịch rảnh", "study_goals, study_plans, student_preferences"],
        ["2", "RAG Chat Tutor", "Gia sư ảo hỏi đáp bám sát tài liệu (Vector Search + Memory)", "chat_sessions, chat_messages, study_material_embeddings"],
        ["3", "Quiz Generator", "Sinh đề thi trắc nghiệm JSON từ tài liệu RAG", "quizzes, study_material_embeddings"],
        ["4", "Recommender", "Đề xuất ôn tập thích ứng khi điểm yếu (HITL Bypass)", "ai_recommendation_reviews, study_plans, notifications"],
        ["5", "Cohort Analytics", "Phân tích lỗ hổng kiến thức tập thể cả lớp", "learning_analytics, classroom_students"],
    ],
    col_widths=[1, 4, 6.5, 6.5]
)

doc.add_heading("4.3. Phân quyền Enum trên bảng users", level=2)
add_para("Enum role: 'student' | 'teacher' | 'admin'", bold=True)
add_para("Enum grade (chỉ bắt buộc khi role = 'student'):", bold=True)
add_table(
    ["Nhóm", "Giá trị Enum"],
    [
        ["Cấp 2 (THCS)", "grade_6, grade_7, grade_8, grade_9"],
        ["Cấp 3 (THPT)", "grade_10, grade_11, grade_12"],
        ["Đại học", "uni_year_1, uni_year_2, uni_year_3, uni_year_4"],
    ],
    col_widths=[5, 13]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. SCOPE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("5. SCOPE", level=1)

doc.add_heading("5.1. Nằm trong phạm vi (In-Scope)", level=2)

add_para("Phân hệ Giáo viên (Teacher Portal) — 4 nhóm CRUD:", bold=True)
add_table(
    ["#", "Nhóm chức năng", "Mã yêu cầu", "Bảng tương tác"],
    [
        ["1", "Quản lý Lớp học", "FR-TCH-CLASS-01 → 04", "classrooms, subjects, classroom_students"],
        ["2", "Quản lý Học sinh trong lớp", "FR-TCH-STU-01 → 03", "classroom_students, users"],
        ["3", "Quản lý Kho tài liệu RAG", "FR-TCH-DOC-01 → 03", "study_documents (MySQL), study_material_embeddings (MongoDB)"],
        ["4", "Dashboard Giám sát Học lực AI", "FR-TCH-DASH-01 → 02", "learning_analytics, ai_recommendation_reviews"],
    ],
    col_widths=[1, 5, 4.5, 7.5]
)

add_para("Phân hệ Học sinh (Student App) — 5 nhóm CRUD:", bold=True)
add_table(
    ["#", "Nhóm chức năng", "Mã yêu cầu", "Bảng tương tác"],
    [
        ["5", "Cấu hình học tập", "FR-STU-PREF-01 → 02", "student_preferences"],
        ["6", "Tham gia Lớp học", "FR-STU-JOIN-01 → 02", "classrooms, classroom_students"],
        ["7", "Lộ trình & Nhiệm vụ hàng ngày", "FR-STU-ROAD-01 → 03", "study_goals, study_plans, study_plan_progress"],
        ["8", "Gia sư ảo RAG + Luyện thi AI", "FR-STU-CHAT-01, FR-STU-QUIZ-01 → 02", "chat_sessions/messages (Mongo), quizzes, quiz_attempts"],
        ["9", "Thông báo & Ôn tập thích ứng", "FR-STU-ALERT-01 → 02", "notifications, study_plans, ai_recommendation_reviews"],
    ],
    col_widths=[1, 5, 5, 7]
)

add_para("Phân hệ Quản trị (Admin Portal) — 3 nhóm CRUD:", bold=True)
add_table(
    ["#", "Nhóm chức năng", "Mã yêu cầu", "Bảng tương tác"],
    [
        ["10", "Quản lý Danh mục Môn học", "FR-ADM-SUB-01 → 04", "subjects"],
        ["11", "Quản lý Tài khoản & Quyền", "FR-ADM-USER-01 → 04", "users"],
        ["12", "Giám sát AI Agent Logs", "FR-ADM-LOG-01 → 02", "ai_logs (MongoDB)"],
    ],
    col_widths=[1, 5.5, 5, 6.5]
)

add_para("Cơ chế đặc biệt:", bold=True)
add_table(
    ["Cơ chế", "Mô tả"],
    [
        ["HITL Bypass", "AI tự duyệt đề xuất ôn tập → status = 'approved' → tạo StudyPlan ngày mai + gửi Notification → không cần GV duyệt"],
        ["Cross-DB Sync (Xóa tài liệu)", "Xóa MySQL study_documents → Xóa Cloudinary file → Xóa MongoDB study_material_embeddings chunks"],
        ["Post-Grading Chain", "Nộp bài → Chấm điểm → Cập nhật study_plans.status → Cập nhật learning_analytics → Nếu yếu: tạo đề xuất AI + lịch bổ sung"],
    ],
    col_widths=[5, 13]
)

doc.add_heading("5.2. Nằm ngoài phạm vi (Out-of-Scope)", level=2)
add_table(
    ["#", "Tính năng", "Lý do", "Giai đoạn"],
    [
        ["1", "Thanh toán học phí (MoMo/VNPAY)", "Chưa xác định mô hình kinh doanh", "Phase 2"],
        ["2", "Video Call trực tuyến (WebRTC)", "Phức tạp hạ tầng", "Phase 2"],
        ["3", "Mobile App Native (iOS/Android)", "Ưu tiên PWA responsive trước", "Phase 3"],
        ["4", "Gamification & Leaderboard", "Tính năng nâng cao", "Phase 2"],
        ["5", "Đề thi tự luận (chấm bằng AI)", "Yêu cầu AI phức tạp hơn", "Phase 3"],
        ["6", "Quản lý chương trình đào tạo", "Ngoài phạm vi MVP", "Phase 2"],
    ],
    col_widths=[1, 6.5, 5.5, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("6. FUNCTIONAL REQUIREMENTS", level=1)

# ── 6.1 TEACHER ──
doc.add_heading("6.1. PHÂN HỆ GIÁO VIÊN (TEACHER WEB PORTAL)", level=2)
add_para(
    "Giáo viên là người làm chủ các lớp học, quản lý học sinh tham gia, cung cấp tài liệu giảng dạy "
    "và theo dõi kết quả học tập thông qua phân tích của AI.",
    italic=True
)

# 6.1.1
doc.add_heading("6.1.1. Quản lý Lớp học (Classroom CRUD)", level=3)

doc.add_heading("FR-TCH-CLASS-01 — Tạo lớp học (Create)", level=4)
add_para("Mô tả: Giáo viên nhập tên lớp, mô tả lớp học và chọn một Môn học (subject_id). "
         "Hệ thống tự động sinh một mã lớp học ngẫu nhiên, duy nhất (class_code, ví dụ: JAVA2026A) "
         "để cung cấp cho học sinh tham gia.")
add_para("Endpoint: POST /api/v1/classrooms", bold=True)
add_para("Luồng xử lý:")
add_bullet("Giáo viên mở form → Frontend nạp dropdown môn học từ API GET /subjects.")
add_bullet("Giáo viên nhập class_name, description, chọn subject_id → gửi request.")
add_bullet("Backend xác thực JWT → lấy teacher_id, validate role == 'teacher'.")
add_bullet("Sinh class_code ngẫu nhiên (8 ký tự, A-Z + 0-9), kiểm tra trùng lặp trong DB.")
add_bullet("INSERT vào bảng classrooms → trả về 201 Created kèm mã class_code.")
add_para("Quy tắc validate:", bold=True)
add_table(
    ["Trường", "Quy tắc", "HTTP nếu lỗi"],
    [
        ["class_name", "Bắt buộc, 3–255 ký tự", "422"],
        ["subject_id", "Phải tồn tại trong bảng subjects", "404"],
        ["description", "Tùy chọn, tối đa 2000 ký tự", "422"],
        ["JWT role", "Phải là teacher", "403"],
    ],
    col_widths=[4, 8, 6]
)
add_para("Bảng tương tác: classrooms, subjects.", italic=True)

doc.add_heading("FR-TCH-CLASS-02 — Xem danh sách & Chi tiết lớp (Read)", level=4)
add_para("Endpoint: GET /api/v1/classrooms", bold=True)
add_para("Mô tả: Xem tất cả lớp giáo viên phụ trách, số lượng học sinh mỗi lớp, thông tin môn học. "
         "Hỗ trợ phân trang (page, page_size).")
add_para("Bảng tương tác: classrooms, subjects, classroom_students.", italic=True)

doc.add_heading("FR-TCH-CLASS-03 — Chỉnh sửa lớp học (Update)", level=4)
add_para("Endpoint: PUT /api/v1/classrooms/{classroom_id}", bold=True)
add_para("Mô tả: Chỉ được cập nhật class_name và description. KHÔNG được đổi subject_id hoặc class_code "
         "(học sinh đã sử dụng mã này để tham gia).")
add_para("Bảo mật: classroom.teacher_id == current_user.id → nếu sai: 403 Forbidden.")
add_para("Bảng tương tác: classrooms.", italic=True)

doc.add_heading("FR-TCH-CLASS-04 — Xóa lớp học (Delete)", level=4)
add_para("Endpoint: DELETE /api/v1/classrooms/{classroom_id}", bold=True)
add_para("Mô tả: Xóa lớp học. FK ON DELETE CASCADE tự động xóa tất cả classroom_students liên quan.")
add_note("Hệ quả cascade: Toàn bộ bản ghi classroom_students bị xóa vĩnh viễn. "
         "Học sinh mất quyền truy cập tài liệu và báo cáo liên quan. Cần hiển thị modal xác nhận trên UI.")
add_para("Bảng tương tác: classrooms, classroom_students.", italic=True)

# 6.1.2
doc.add_heading("6.1.2. Quản lý Học sinh trong lớp (Student Membership)", level=3)

doc.add_heading("FR-TCH-STU-01 — Thêm học sinh vào lớp (Create)", level=4)
add_para("Endpoint: POST /api/v1/classrooms/{classroom_id}/students", bold=True)
add_para("Mô tả: Giáo viên thêm học sinh bằng email. Hệ thống tìm ID học sinh trong bảng users "
         "và tạo liên kết trong bảng classroom_students.")
add_para("Luồng xử lý:")
add_bullet("Tìm user theo email → nếu không tồn tại: 404.")
add_bullet("Kiểm tra role == 'student' → nếu không: 400.")
add_bullet("Kiểm tra chưa tham gia lớp (classroom_students trùng) → nếu trùng: 409.")
add_bullet("INSERT classroom_students → 201 Created.")
add_para("Bảng tương tác: classroom_students, users, classrooms.", italic=True)

doc.add_heading("FR-TCH-STU-02 — Xem danh sách học sinh trong lớp (Read)", level=4)
add_para("Endpoint: GET /api/v1/classrooms/{classroom_id}/students", bold=True)
add_para("Mô tả: Hiển thị Họ tên, Email, Ảnh đại diện, Khối lớp (grade), Ngày gia nhập (joined_at).")
add_para("Bảng tương tác: classroom_students, users.", italic=True)

doc.add_heading("FR-TCH-STU-03 — Xóa học sinh khỏi lớp (Delete)", level=4)
add_para("Endpoint: DELETE /api/v1/classrooms/{classroom_id}/students/{student_id}", bold=True)
add_para("Mô tả: Trục xuất học sinh khỏi lớp. Xóa bản ghi classroom_students. "
         "KHÔNG xóa tài khoản học sinh, chỉ xóa liên kết lớp.")
add_para("Lưu ý: Quyền truy cập tài liệu gắn theo subject_id (qua study_goals), "
         "không gắn theo classroom_id. Nếu HS có study_goals active cho môn đó, vẫn truy cập được RAG.",
         italic=True)

# 6.1.3
doc.add_heading("6.1.3. Quản lý Kho tài liệu học tập (Study Document CRUD)", level=3)

doc.add_heading("FR-TCH-DOC-01 — Tải lên tài liệu mới (Create) — Đồng bộ AI", level=4)
add_para("Endpoint: POST /api/v1/study-documents (multipart/form-data)", bold=True)
add_para("Mô tả: Giáo viên chọn môn học, nhập tiêu đề, chọn file (PDF/DOCX/TXT/MD). "
         "Backend xử lý đồng bộ với AI Vector Store.")
add_para("Luồng xử lý Cross-Database:", bold=True)
add_bullet("Bước 1 — Lưu metadata MySQL: Upload file lên Cloudinary → lấy file_path URL → "
           "INSERT INTO study_documents (subject_id, created_by, title, file_path, file_type).")
add_bullet("Bước 2 — Background Task Vector hóa: Đọc nội dung file → chia nhỏ thành chunks "
           "(500–1000 tokens, overlap 100 tokens) → sinh embedding 4096 chiều (NVIDIA NIM) → "
           "INSERT study_material_embeddings (MongoDB).")
add_para("Quy tắc validate:", bold=True)
add_table(
    ["Trường", "Quy tắc", "HTTP"],
    [
        ["file", "Chỉ .pdf, .docx, .txt, .md", "415"],
        ["Kích thước", "Tối đa 50 MB", "413"],
        ["subject_id", "Phải tồn tại trong subjects", "404"],
        ["JWT role", "Phải là teacher", "403"],
    ],
    col_widths=[4, 10, 4]
)
add_para("Bảng tương tác: study_documents (MySQL), study_material_embeddings (MongoDB), Cloudinary.", italic=True)

doc.add_heading("FR-TCH-DOC-02 — Xem kho tài liệu (Read)", level=4)
add_para("Endpoint: GET /api/v1/study-documents?subject_id=3", bold=True)
add_para("Mô tả: Bộ lọc theo Môn học. Hiển thị tên file, file_type, ngày tải, link tải gốc.")
add_para("Bảng tương tác: study_documents, subjects.", italic=True)

doc.add_heading("FR-TCH-DOC-03 — Xóa tài liệu (Delete) — Phức tạp nhất", level=4)
add_para("Endpoint: DELETE /api/v1/study-documents/{document_id}", bold=True)
add_para("Mô tả: Xóa đồng bộ ở 3 nơi: MySQL + Cloudinary + MongoDB Vector Store.")
add_para("Luồng Đồng bộ Xóa (Cross-Database Cascade):", bold=True)
add_bullet("Bước 1: Kiểm tra quyền sở hữu (created_by == current_user.id).")
add_bullet("Bước 2: DELETE file từ Cloudinary.")
add_bullet("Bước 3: DELETE toàn bộ vector chunks trong MongoDB study_material_embeddings "
           "(WHERE metadata.document_id = document_id).")
add_bullet("Bước 4: DELETE bản ghi trong MySQL study_documents.")
add_note("Toàn vẹn dữ liệu cross-database: Nếu MongoDB xóa lỗi nhưng MySQL đã xóa → AI có 'dữ liệu rác' "
         "(orphaned vectors). Giải pháp: try-catch wrap, nếu MongoDB fail → rollback MySQL transaction.")
add_para("Bảng tương tác: study_documents (MySQL), study_material_embeddings (MongoDB), Cloudinary.", italic=True)

# 6.1.4
doc.add_heading("6.1.4. Dashboard Giám sát Học lực & Ôn tập AI", level=3)

doc.add_heading("FR-TCH-DASH-01 — Báo cáo học tập cá nhân (Read)", level=4)
add_para("Endpoint: GET /api/v1/classrooms/{classroom_id}/analytics", bold=True)
add_para("Mô tả: Xem hồ sơ learning_analytics từng học sinh: điểm trung bình, số đề đã làm, "
         "danh sách chủ đề mạnh (strong_topics), chủ đề yếu (weak_topics), nhận xét AI (ai_feedback). "
         "Bao gồm thống kê tổng hợp cả lớp: class_average_score, common_weak_topics.")
add_para("Bảng tương tác: learning_analytics, classroom_students, users.", italic=True)

doc.add_heading("FR-TCH-DASH-02 — Theo dõi đề xuất ôn tập AI + Viết phản hồi (Read + Update)", level=4)
add_para("Endpoint đọc: GET /api/v1/classrooms/{classroom_id}/recommendations", bold=True)
add_para("Endpoint ghi: PATCH /api/v1/recommendations/{review_id}/feedback", bold=True)
add_para("Mô tả: Giáo viên xem đề xuất AI đã tự duyệt (status = 'approved'). "
         "Giáo viên có thể viết thêm teacher_feedback để bổ sung bài tập hoặc dặn dò thêm.")
add_para("Bảng tương tác: ai_recommendation_reviews, users, classroom_students.", italic=True)

doc.add_page_break()

# ── 6.2 STUDENT ──
doc.add_heading("6.2. PHÂN HỆ HỌC SINH (STUDENT MOBILE/WEB APP)", level=2)
add_para(
    "Học sinh tham gia lớp học, tự đặt mục tiêu, tự học cùng Gia sư AI và luyện đề kiểm tra "
    "để tích lũy tiến độ hoàn thành lộ trình học tập.",
    italic=True
)

# 6.2.1
doc.add_heading("6.2.1. Cấu hình học tập (Student Preferences)", level=3)

doc.add_heading("FR-STU-PREF-01 — Xem cấu hình hiện tại (Read)", level=4)
add_para("Endpoint: GET /api/v1/students/me/preferences", bold=True)
add_para("Mô tả: Xem số giờ cam kết học mỗi ngày, buổi học ưu tiên và lịch rảnh chi tiết 7 ngày/tuần.")
add_para("Bảng tương tác: student_preferences.", italic=True)

doc.add_heading("FR-STU-PREF-02 — Cập nhật cấu hình (Upsert)", level=4)
add_para("Endpoint: PUT /api/v1/students/me/preferences", bold=True)
add_para("Mô tả: Thay đổi số giờ cam kết, chọn lại buổi ưu tiên (Sáng/Chiều/Tối), bật/tắt khung giờ rảnh. "
         "Cấu hình mới ngay lập tức được áp dụng cho các lộ trình AI sinh ra sau đó.")
add_para("Quy tắc validate:", bold=True)
add_table(
    ["Trường", "Quy tắc"],
    [
        ["study_hours_per_day", "Integer, 1 ≤ giá trị ≤ 12"],
        ["preferred_study_time", "Enum: 'morning' | 'afternoon' | 'evening'"],
        ["available_schedule", "JSON Object — key phải thuộc {mon, tue, wed, thu, fri, sat, sun}"],
    ],
    col_widths=[6, 12]
)
add_para("Mối liên hệ AI: Goal Planner Agent đọc student_preferences để chỉ phân bổ nhiệm vụ vào khung giờ rảnh.",
         italic=True)
add_para("Bảng tương tác: student_preferences.", italic=True)

# 6.2.2
doc.add_heading("6.2.2. Đăng ký tham gia Lớp học", level=3)

doc.add_heading("FR-STU-JOIN-01 — Tham gia lớp bằng mã code (Create)", level=4)
add_para("Endpoint: POST /api/v1/students/me/join-classroom", bold=True)
add_para("Mô tả: Học sinh nhập class_code → hệ thống tìm lớp → kiểm tra chưa tham gia → INSERT classroom_students.")
add_para("Quy tắc: Mã không tồn tại → 404. Đã tham gia → 409.")
add_para("Bảng tương tác: classrooms, classroom_students.", italic=True)

doc.add_heading("FR-STU-JOIN-02 — Rời lớp học (Delete)", level=4)
add_para("Endpoint: DELETE /api/v1/students/me/classrooms/{classroom_id}", bold=True)
add_para("Mô tả: Xóa classroom_students, giữ nguyên toàn bộ dữ liệu cá nhân (study_goals, quiz_attempts...).")
add_para("Bảng tương tác: classroom_students.", italic=True)

# 6.2.3
doc.add_heading("6.2.3. Lộ trình & Nhiệm vụ hàng ngày (Roadmap)", level=3)

doc.add_heading("FR-STU-ROAD-01 — Tạo lộ trình học tập mới (Create) — AI Auto-Scheduling", level=4)
add_para("Endpoint: POST /api/v1/students/me/study-goals", bold=True)
add_para("Mô tả: Chọn môn, nhập điểm mục tiêu + deadline. AI tự động lập lịch học hàng ngày.")
add_para("Luồng xử lý:", bold=True)
add_bullet("Bước 1: Validate subject_id + kiểm tra không trùng goal active → INSERT study_goals (status='active').")
add_bullet("Bước 2: Đọc student_preferences (lịch rảnh, giờ ưu tiên, số giờ cam kết).")
add_bullet("Bước 3: Gọi Goal Planner Agent → AI sinh mảng nhiệm vụ [{title, description, study_date, start_time, end_time}].")
add_bullet("Bước 4: INSERT study_plans (ai_generated=True, status='todo') cho mỗi nhiệm vụ.")
add_para("Quy tắc validate:", bold=True)
add_table(
    ["Trường", "Quy tắc", "HTTP"],
    [
        ["target_score", "1.0 ≤ giá trị ≤ 10.0, DECIMAL(4,2)", "422"],
        ["deadline", "Phải > ngày hiện tại (tối thiểu T+1)", "422"],
        ["Không trùng goal active", "Kiểm tra study_goals cùng student_id + subject_id + status='active'", "409"],
        ["subject_id", "Phải tồn tại trong subjects", "404"],
    ],
    col_widths=[5, 9, 4]
)
add_para("Bảng tương tác: study_goals, study_plans, student_preferences, subjects.", italic=True)

doc.add_heading("FR-STU-ROAD-02 — Cập nhật tiến độ nhiệm vụ ngày (Update)", level=4)
add_para("Endpoint: PATCH /api/v1/students/me/study-plans/{plan_id}/progress", bold=True)
add_para("2 cơ chế cập nhật:", bold=True)
add_table(
    ["Cơ chế", "Điều kiện", "Kết quả"],
    [
        ["A — Tự động (từ bài thi)", "quiz_attempts.score ≥ 5.0 (quiz gắn study_plan_id)", "study_plans.status → 'done'"],
        ["B — Thủ công (đọc tài liệu)", "HS cập nhật completion_percent (0–100%)", "≥100% → 'done', >0% → 'doing'"],
    ],
    col_widths=[5, 7, 6]
)
add_para("Bảng trạng thái study_plans.status:", bold=True)
add_table(
    ["Trạng thái", "Điều kiện chuyển"],
    [
        ["todo", "Mặc định khi AI sinh"],
        ["doing", "completion_percent > 0 AND < 100"],
        ["done", "completion_percent ≥ 100 HOẶC quiz.score ≥ 5.0"],
    ],
    col_widths=[4, 14]
)
add_para("Bảng tương tác: study_plans, study_plan_progress.", italic=True)

doc.add_heading("FR-STU-ROAD-03 — Dừng / Chỉnh sửa lộ trình (Update/Delete)", level=4)
add_para("Endpoint hủy: PATCH /api/v1/students/me/study-goals/{goal_id}/drop", bold=True)
add_para("Endpoint sửa: PUT /api/v1/students/me/study-goals/{goal_id}", bold=True)
add_para("Hủy: Chuyển study_goals.status = 'dropped' + xóa study_plans WHERE status IN ('todo', 'doing').")
add_para("Sửa mục tiêu: Cập nhật target_score/deadline → xóa plans chưa xong → AI lập lại lịch mới.")
add_para("Bảng tương tác: study_goals, study_plans.", italic=True)

# 6.2.4
doc.add_heading("6.2.4. Gia sư ảo & Luyện thi AI (RAG Study & AI Quiz)", level=3)

doc.add_heading("FR-STU-CHAT-01 — Chat hỏi bài Gia sư ảo (RAG Chat + SSE)", level=4)
add_para("Endpoint tạo phiên: POST /api/v1/students/me/chat/sessions", bold=True)
add_para("Endpoint gửi tin: POST /api/v1/students/me/chat/sessions/{id}/messages (SSE Stream)", bold=True)
add_para("Luồng RAG:", bold=True)
add_bullet("Bước 1: Lưu tin nhắn user vào MongoDB chat_messages + đọc 20 tin gần nhất (conversation memory).")
add_bullet("Bước 2: RAG Vector Search — tìm top 5 chunks liên quan trong study_material_embeddings (cosine similarity).")
add_bullet("Bước 3: Xây dựng prompt = System prompt + 5 chunks tài liệu + conversation history + câu hỏi.")
add_bullet("Bước 4: Gọi OpenAI API (stream=True) → truyền từng token về Frontend qua SSE real-time.")
add_bullet("Bước 5: Lưu full response vào chat_messages (role: assistant) + ghi ai_logs.")
add_para("Cơ chế Memory hội thoại:", bold=True)
add_table(
    ["Loại Memory", "Khi nào áp dụng", "Cách hoạt động"],
    [
        ["ConversationBufferMemory", "Phiên < 20 tin nhắn", "Lưu nguyên văn 20 tin gần nhất"],
        ["SummaryMemory", "Phiên > 20 tin nhắn", "Tóm tắt tin cũ thành 1 đoạn ngắn"],
    ],
    col_widths=[5.5, 5, 7.5]
)
add_para("Bảng tương tác: chat_sessions, chat_messages, study_material_embeddings (MongoDB).", italic=True)

doc.add_heading("FR-STU-QUIZ-01 — Tự tạo Quiz kiểm tra (AI Quiz Generation)", level=4)
add_para("Endpoint: POST /api/v1/students/me/quizzes/generate", bold=True)
add_para("Mô tả: Chọn bài học trong lộ trình → AI sinh đề thi trắc nghiệm JSON bám sát tài liệu RAG. "
         "Đề thi gồm câu hỏi, 4 phương án (A/B/C/D), đáp án đúng, và giải thích chi tiết.")
add_para("Tham số đầu vào: study_plan_id (bài học trong lộ trình), difficulty (easy/medium/hard), total_questions.")
add_para("Bảng tương tác: quizzes, study_plans (MySQL), study_material_embeddings (MongoDB).", italic=True)

doc.add_heading("FR-STU-QUIZ-02 — Nộp bài & Auto-Grade + Post-Grading Chain", level=4)
add_para("Endpoint: POST /api/v1/students/me/quizzes/{quiz_id}/submit", bold=True)
add_para("Mô tả: Nộp bài → chấm tự động → kích hoạt chuỗi nghiệp vụ hậu chấm thi.", bold=True)
add_para("Chuỗi nghiệp vụ hậu chấm thi (Post-Grading Chain):", bold=True)
add_table(
    ["Bước", "Điều kiện", "Hành động"],
    [
        ["1. Chấm điểm", "Luôn luôn", "So sánh answers với correct_answer JSON → tính score → INSERT quiz_attempts"],
        ["2. Cập nhật plan", "score ≥ 5.0 AND quiz có study_plan_id", "UPDATE study_plans SET status = 'done'"],
        ["3. Cập nhật analytics", "Luôn luôn", "UPDATE learning_analytics (average_score, weak/strong topics, ai_feedback)"],
        ["4. HITL Bypass", "score < 8.0", "INSERT ai_recommendation_reviews (status='approved') + INSERT study_plans ngày mai (19:00–20:00) + INSERT notifications"],
        ["5. Adaptive Plan", "score < 7.0", "AI điều chỉnh lại lộ trình tương lai + gửi notification"],
    ],
    col_widths=[3.5, 4, 10.5]
)
add_note("HITL Bypass: AI tự duyệt đề xuất ôn tập (status = 'approved') và tạo lịch học bổ sung ngày mai "
         "mà không cần giáo viên duyệt thủ công. Đây là cơ chế cốt lõi để loại bỏ nghẽn cổ chai.")
add_para("Bảng tương tác: quiz_attempts, quizzes, study_plans, learning_analytics, "
         "ai_recommendation_reviews, notifications.", italic=True)

# 6.2.5
doc.add_heading("6.2.5. Thông báo & Ôn tập thích ứng (Adaptive Alerts)", level=3)

doc.add_heading("FR-STU-ALERT-01 — Nhận lịch học bổ sung tự động", level=4)
add_para("Mô tả: Nhiệm vụ có prefix [Ôn tập AI] tự xuất hiện ngày mai nếu điểm hôm nay < 8.0. "
         "study_date = today+1, start_time = 19:00, end_time = 20:00, ai_generated = True.")
add_para("Bảng tương tác: study_plans, ai_recommendation_reviews.", italic=True)

doc.add_heading("FR-STU-ALERT-02 — Xem & Quản lý thông báo (CRUD)", level=4)
add_para("Endpoint đọc: GET /api/v1/students/me/notifications", bold=True)
add_para("Endpoint đánh dấu đã đọc: PATCH /api/v1/notifications/{id}/read", bold=True)
add_para("Endpoint xóa: DELETE /api/v1/notifications/{id}", bold=True)
add_para("Các loại notifications.type:", bold=True)
add_table(
    ["Enum", "Nguồn phát sinh", "Ví dụ"],
    [
        ["system", "Admin bảo trì", "Hệ thống bảo trì lúc 02:00 AM"],
        ["plan", "AI tạo lịch bổ sung", "Lịch ôn tập AI bổ sung đã được tạo"],
        ["recommendation", "AI đề xuất ôn tập", "AI gợi ý ôn lại chương Đa luồng"],
        ["quiz", "Kết quả chấm thi", "Bài thi Đa luồng: 6.5/10"],
        ["score", "Cập nhật hồ sơ", "Hồ sơ học tập được cập nhật"],
    ],
    col_widths=[4, 4, 10]
)
add_para("Bảng tương tác: notifications.", italic=True)

doc.add_page_break()

# ── 6.3 ADMIN ──
doc.add_heading("6.3. PHÂN HỆ QUẢN TRỊ VIÊN (SYSTEM ADMIN PORTAL)", level=2)

# 6.3.1
doc.add_heading("6.3.1. Quản lý Danh mục Môn học (Subject CRUD)", level=3)

doc.add_heading("FR-ADM-SUB-01 — Tạo môn học (Create)", level=4)
add_para("Endpoint: POST /api/v1/admin/subjects", bold=True)
add_para("Quy tắc: name (2–255 ký tự), code (UNIQUE, A-Z 0-9 underscore), description (tùy chọn, max 5000).")

doc.add_heading("FR-ADM-SUB-02 — Xem danh sách (Read)", level=4)
add_para("Endpoint: GET /api/v1/admin/subjects?search=java", bold=True)
add_para("Hỗ trợ tìm kiếm theo tên hoặc mã môn học, phân trang.")

doc.add_heading("FR-ADM-SUB-03 — Cập nhật (Update)", level=4)
add_para("Endpoint: PUT /api/v1/admin/subjects/{id}", bold=True)
add_para("Cho phép đổi name, description. KHÔNG cho phép đổi code sau khi có lớp/tài liệu liên kết.")

doc.add_heading("FR-ADM-SUB-04 — Xóa môn học (Delete)", level=4)
add_para("Endpoint: DELETE /api/v1/admin/subjects/{id}", bold=True)
add_note("Ràng buộc xóa an toàn: Chỉ xóa khi KHÔNG có bản ghi nào trong classrooms, study_documents, "
         "study_goals, quizzes, learning_analytics tham chiếu đến subject_id này. Nếu còn → 409 Conflict.")
add_para("Bảng tương tác: subjects.", italic=True)

# 6.3.2
doc.add_heading("6.3.2. Quản lý Tài khoản & Quyền hạn (User CRUD)", level=3)

doc.add_heading("FR-ADM-USER-01 — Tạo tài khoản (Create)", level=4)
add_para("Endpoint: POST /api/v1/admin/users", bold=True)
add_para("Quy tắc: email (UNIQUE), password (≥ 8, có 1 chữ hoa + 1 số), role (Enum), "
         "grade (bắt buộc nếu student, NULL nếu teacher/admin).")

doc.add_heading("FR-ADM-USER-02 — Xem danh sách (Read)", level=4)
add_para("Endpoint: GET /api/v1/admin/users?role=student&search=nguyen&is_active=true", bold=True)
add_para("Hỗ trợ lọc theo role, is_active, tìm kiếm theo tên/email, phân trang.")

doc.add_heading("FR-ADM-USER-03 — Cập nhật / Khóa tài khoản (Update)", level=4)
add_para("Endpoint: PUT /api/v1/admin/users/{id}", bold=True)
add_para("Trường cho phép: full_name, role, grade, is_active (False = khóa), password (reset).")
add_note("Khóa tài khoản (is_active = False): JWT hiện tại vẫn hoạt động đến khi hết hạn. "
         "Cần kiểm tra is_active ở middleware mỗi request HOẶC invalidate JWT trong Redis.")

doc.add_heading("FR-ADM-USER-04 — Xóa tài khoản vĩnh viễn (Delete)", level=4)
add_para("Endpoint: DELETE /api/v1/admin/users/{id}", bold=True)
add_note("CASCADE xóa 12 bảng liên quan: student_preferences, classroom_students, classrooms "
         "(nếu teacher), study_goals, study_plans, study_plan_progress, study_documents, quizzes, "
         "quiz_attempts, learning_analytics, ai_recommendation_reviews, notifications. "
         "Xóa giáo viên → cascade xóa TOÀN BỘ lớp học. Cần hiển thị cảnh báo chi tiết.")
add_para("Bảng tương tác: users + 12 bảng cascade.", italic=True)

# 6.3.3
doc.add_heading("6.3.3. Giám sát AI Agent Logs", level=3)

doc.add_heading("FR-ADM-LOG-01 — Xem lịch sử gọi LLM (Read)", level=4)
add_para("Endpoint: GET /api/v1/admin/ai-logs?agent_name=QuizGenerator", bold=True)
add_para("Mô tả: Hiển thị agent_name, student_id, input, output, execution_time, created_at. "
         "Kèm thống kê tổng cuộc gọi/ngày, thời gian trung bình, breakdown theo agent.")

doc.add_heading("FR-ADM-LOG-02 — Xóa log cũ (Cleanup)", level=4)
add_para("Endpoint: DELETE /api/v1/admin/ai-logs/cleanup?older_than_days=30", bold=True)
add_para("Mô tả: Xóa các log đã quá 30 ngày trong MongoDB ai_logs để tiết kiệm dung lượng.")
add_para("Collection tương tác: ai_logs (MongoDB).", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("7. NON-FUNCTIONAL REQUIREMENTS", level=1)

doc.add_heading("7.1. Hiệu năng & Khả năng mở rộng", level=2)
add_table(
    ["Mã", "Yêu cầu", "Chỉ số"],
    [
        ["NFR-PERF-01", "AI Chat Tutor phản hồi First Token nhanh", "≤ 2.5 giây (SSE streaming)"],
        ["NFR-PERF-02", "API CRUD thông thường phản hồi nhanh", "≤ 500ms (P95)"],
        ["NFR-PERF-03", "AI sinh đề thi 10 câu", "≤ 30 giây"],
        ["NFR-PERF-04", "Redis cache giảm tải MySQL", "Giảm 40% truy vấn MySQL"],
        ["NFR-PERF-05", "Tránh tràn ID khóa chính", "Toàn bộ PK MySQL dùng BIGINT (max 9.2×10¹⁸)"],
    ],
    col_widths=[3.5, 7, 7.5]
)

doc.add_heading("7.2. Bảo mật", level=2)
add_table(
    ["Mã", "Yêu cầu", "Chi tiết"],
    [
        ["NFR-SEC-01", "JWT Token bảo mật", "Lưu Cookie HTTP-Only (chống XSS). Thời hạn 24 giờ"],
        ["NFR-SEC-02", "Mật khẩu mã hóa", "bcrypt hash. Không lưu plaintext"],
        ["NFR-SEC-03", "Phân quyền API", "Middleware kiểm tra role + is_active mỗi request"],
        ["NFR-SEC-04", "Kiểm tra quyền sở hữu", "Teacher chỉ CRUD lớp CỦA MÌNH. Student chỉ truy cập CỦA MÌNH"],
    ],
    col_widths=[3.5, 5, 9.5]
)

doc.add_heading("7.3. Dữ liệu & Đồng bộ", level=2)
add_table(
    ["Mã", "Yêu cầu", "Chi tiết"],
    [
        ["NFR-DATA-01", "Cross-DB sync khi xóa tài liệu", "MySQL DELETE + Cloudinary DELETE + MongoDB deleteMany. Rollback nếu MongoDB fail"],
        ["NFR-DATA-02", "Liên kết cross-database", "Qua student_id/subject_id (Int) — tham chiếu logic, KHÔNG FK vật lý"],
        ["NFR-DATA-03", "Timestamp UTC", "Mọi DATETIME lưu UTC. Frontend chuyển timezone local"],
        ["NFR-DATA-04", "Enum grade chỉ bắt buộc student", "teacher/admin → grade = NULL"],
    ],
    col_widths=[3.5, 5, 9.5]
)

doc.add_heading("7.4. AI Agent", level=2)
add_table(
    ["Mã", "Yêu cầu", "Chi tiết"],
    [
        ["NFR-AI-01", "HITL Bypass mặc định", "ai_recommendation_reviews.status = 'approved' tự động"],
        ["NFR-AI-02", "Logging bắt buộc", "Mọi cuộc gọi LLM ghi vào MongoDB ai_logs"],
        ["NFR-AI-03", "Goal Planner đọc preferences", "Không sinh study_plans vào giờ HS không rảnh"],
        ["NFR-AI-04", "Adaptive threshold", "< 8.0 → đề xuất ôn tập. < 7.0 → điều chỉnh lộ trình"],
    ],
    col_widths=[3.5, 5.5, 9]
)

doc.add_heading("7.5. Hạ tầng & Deployment", level=2)
add_table(
    ["Mã", "Yêu cầu", "Chi tiết"],
    [
        ["NFR-INFRA-01", "Container hóa toàn bộ", "Docker + docker-compose: FastAPI, Next.js, MySQL, MongoDB, Redis, Nginx"],
        ["NFR-INFRA-02", "Reverse proxy", "Nginx xử lý SSL, routing, static files"],
        ["NFR-INFRA-03", "File Storage", "Cloudinary (production) / Local (dev fallback)"],
    ],
    col_widths=[3.5, 5, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. SUCCESS METRICS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("8. SUCCESS METRICS", level=1)

doc.add_heading("8.1. Nhóm 1: Engagement (Mức độ tương tác)", level=2)
add_table(
    ["Metric", "Công thức đo", "Mục tiêu", "Nguồn dữ liệu"],
    [
        ["Tỷ lệ hoàn thành lộ trình", "COUNT(status='done') / COUNT(*) × 100%", "≥ 70%", "study_plans"],
        ["Số phiên chat AI / HS / tuần", "COUNT(chat_sessions) GROUP BY student, WEEK", "≥ 3 phiên/tuần", "chat_sessions (MongoDB)"],
        ["Số đề thi đã làm / HS / tuần", "COUNT(quiz_attempts) GROUP BY student, WEEK", "≥ 2 đề/tuần", "quiz_attempts"],
        ["Tỷ lệ cập nhật preferences", "COUNT(student_preferences) / COUNT(students)", "≥ 80%", "student_preferences, users"],
    ],
    col_widths=[4, 6, 3.5, 4.5]
)

doc.add_heading("8.2. Nhóm 2: AI Quality (Chất lượng AI)", level=2)
add_table(
    ["Metric", "Công thức đo", "Mục tiêu", "Nguồn dữ liệu"],
    [
        ["Thời gian First Token", "AVG(execution_time) WHERE agent='RAGTutor'", "≤ 2.5 giây", "ai_logs (MongoDB)"],
        ["Thời gian sinh đề 10 câu", "AVG(execution_time) WHERE agent='QuizGenerator'", "≤ 30 giây", "ai_logs (MongoDB)"],
        ["% đề xuất ôn tập được áp dụng", "COUNT(status='approved') / COUNT(*)", "100%", "ai_recommendation_reviews"],
        ["Điểm TB sau ôn tập AI", "So sánh AVG(score) trước/sau [Ôn tập AI]", "Tăng ≥ 1.0 điểm", "quiz_attempts, study_plans"],
    ],
    col_widths=[4, 6, 3.5, 4.5]
)

doc.add_heading("8.3. Nhóm 3: Performance (Hiệu năng hệ thống)", level=2)
add_table(
    ["Metric", "Công thức đo", "Mục tiêu", "Công cụ"],
    [
        ["API Response Time (P95)", "95th percentile latency CRUD APIs", "≤ 500ms", "Application metrics"],
        ["Uptime", "(Total - Downtime) / Total × 100%", "≥ 99.5%", "Docker health checks"],
        ["Vector Search latency", "AVG($vectorSearch execution time)", "≤ 800ms", "MongoDB profiler"],
    ],
    col_widths=[4.5, 5.5, 4, 4]
)

doc.add_heading("8.4. Nhóm 4: Business Impact (Tác động kinh doanh)", level=2)
add_table(
    ["Metric", "Công thức đo", "Mục tiêu", "Nguồn"],
    [
        ["Thời gian GV soạn đề", "Thủ công (2–4h) vs AI (30s)", "Giảm 99%", "User interview"],
        ["% lớp có báo cáo lỗ hổng tập thể", "COUNT(lớp có common_weak_topics) / COUNT(lớp)", "100%", "learning_analytics"],
        ["Số tài liệu vector hóa", "COUNT(study_material_embeddings)", "≥ 1000 chunks/môn", "MongoDB"],
        ["Tỷ lệ HS cải thiện điểm", "COUNT(HS có avg_score tăng sau 30 ngày) / COUNT(HS)", "≥ 60%", "learning_analytics"],
    ],
    col_widths=[4, 6, 3.5, 4.5]
)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "..", "BRD_EDU_MIND.docx")
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f"✅ Đã tạo file BRD thành công: {output_path}")
