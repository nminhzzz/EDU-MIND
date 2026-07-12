"""
Báo cáo siêu chi tiết: Luồng Luyện thi + AI Recommendation.
Run: py -3 Backend/scripts/generate_detailed_flow_report.py
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = (
    Path(__file__).resolve().parent.parent
    / "materials"
    / "Bao_cao_Chi_tiet_Luyen_thi_va_AI_Recommendation.docx"
)


def font(run, *, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color


def title(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, bold=True, size=size)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    font(r, bold=bold)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for r in para.runs:
        font(r)


def num(doc, text):
    para = doc.add_paragraph(text, style="List Number")
    for r in para.runs:
        font(r)


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].text = hd
        for para in t.rows[0].cells[i].paragraphs:
            for r in para.runs:
                font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for para in t.rows[ri + 1].cells[ci].paragraphs:
                for r in para.runs:
                    font(r, size=10)
    doc.add_paragraph()


def code(doc, text):
    para = doc.add_paragraph()
    r = para.add_run(text)
    font(r, size=10)
    r.font.name = "Consolas"


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)

    # ── TRANG BÌA ───────────────────────────────────────────────────────────
    title(doc, "BÁO CÁO KỸ THUẬT SIÊU CHI TIẾT")
    title(doc, "Luồng Luyện thi trắc nghiệm AI\n& Khuyến nghị học tập AI (HITL)", 14)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        para.add_run(
            f"Dự án: AI Learning Assistant Platform\n"
            f"Phiên bản báo cáo: 2.0 — Chi tiết đầy đủ Frontend + Backend + Agent + DB\n"
            f"Ngày lập: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ),
        size=12,
    )
    page_break(doc)

    # ── MỤC LỤC ─────────────────────────────────────────────────────────────
    h(doc, "MỤC LỤC")
    toc = [
        "PHẦN A — TỔNG QUAN KIẾN TRÚC",
        "PHẦN B — LUỒNG LUYỆN THI (QUIZ)",
        "  B1. Hai nhánh nghiệp vụ: Luyện đề tự do vs Nhiệm vụ ngày",
        "  B2. Luồng Frontend (Student)",
        "  B3. API Reference — Quizzes",
        "  B4. Luồng Backend sinh đề (Generate) — từng bước",
        "  B5. RAG — Vector Search tài liệu",
        "  B6. Multi-Agent QC (Agent-to-Agent)",
        "  B7. Chuẩn hóa câu hỏi & Lưu DB",
        "  B8. Luồng làm bài & Submit",
        "  B9. Cơ chế chấm điểm",
        "  B10. Cấu trúc dữ liệu JSON",
        "  B11. Bảng MySQL / MongoDB liên quan",
        "  B12. Xử lý lỗi & Timeout",
        "PHẦN C — LUỒNG AI RECOMMENDATION",
        "  C1. Điều kiện kích hoạt & Tách sync/async",
        "  C2. Background Task chi tiết",
        "  C3. Analytics Agent",
        "  C4. Recommender Agent",
        "  C5. Human-in-the-Loop — Giáo viên duyệt",
        "  C6. Học sinh nhận kết quả",
        "  C7. Adaptive Plan (score < 7)",
        "  C8. Notification, Email, Cache",
        "  C9. API Reference — Recommendations",
        "  C10. Cấu trúc dữ liệu & Bảng DB",
        "PHẦN D — LUỒNG END-TO-END (50 bước)",
        "PHẦN E — CHỈ MỤC FILE MÃ NGUỒN",
        "PHẦN F — BẢNG NGƯỠNG NGHIỆP VỤ",
        "PHẦN G — KẾT LUẬN",
    ]
    for line in toc:
        p(doc, line)
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN A
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN A — TỔNG QUAN KIẾN TRÚC", 1)
    p(
        doc,
        "Hệ thống gồm hai luồng nghiệp vụ lớn được mô tả chi tiết trong báo cáo này. "
        "Luồng Luyện thi xử lý sinh đề AI (RAG + Multi-Agent QC), làm bài và chấm điểm đồng bộ. "
        "Luồng AI Recommendation chạy ngầm sau submit, phân tích học lực và sinh đề xuất ôn tập "
        "theo mô hình Human-in-the-Loop (HITL) — giáo viên duyệt trước khi học sinh nhận.",
    )
    tbl(
        doc,
        ["Tầng", "Công nghệ", "Vai trò"],
        [
            ["Frontend", "Next.js (App Router)", "UI học sinh / giáo viên"],
            ["API", "FastAPI", "REST endpoints, auth JWT, BackgroundTasks"],
            ["Service", "Python modules", "Nghiệp vụ: quiz/, analytic_service, recommendation_service"],
            ["Agent", "NVIDIA NIM (LLM)", "Sinh đề, QC, Analytics, Recommender"],
            ["MySQL", "SQLAlchemy", "quizzes, attempts, analytics, recommendations"],
            ["MongoDB", "Motor/PyMongo", "Embedding tài liệu, chat sessions"],
            ["Redis", "aioredis", "Cache RAG, cache recommendations"],
        ],
    )
    p(doc, "Mô hình Multi-Agent trong hệ thống:", bold=True)
    bullet(doc, "Sinh đề: Quiz Generator Agent ↔ QC Reviewer Agent (pipeline orchestrated bởi service).")
    bullet(doc, "Sau submit: Analytics Agent → (nếu score<8) Recommender Agent → GV duyệt (HITL).")
    bullet(doc, "Score<7: Roadmap Planner qua chat session (Adaptive Plan).")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN B — LUYỆN THI
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN B — LUỒNG LUYỆN THI (QUIZ)", 1)

    h(doc, "B1. Hai nhánh nghiệp vụ", 2)
    tbl(
        doc,
        ["Tiêu chí", "Luyện đề tự do", "Nhiệm vụ ngày (Study Plan)"],
        [
            ["Kích hoạt", "HS bấm 'Sinh đề AI' tại /student/quizzes", "Tab Quiz trong Task Study"],
            ["study_plan_id", "null", "Có giá trị — liên kết nhiệm vụ"],
            ["QC Review", "BẬT (Multi-Agent)", "TẮT (skip_qc=True) — sinh nhanh"],
            ["Số câu mặc định FE", "5 (cấu hình form)", "5 (quick-quiz-pane)"],
            ["Timeout FE generate", "120 giây", "120 giây"],
            ["Hoàn thành task", "Không", "score ≥ 8 → mark study_plan done"],
            ["Tiêu đề hiển thị lịch sử", "Tên đề / 'Đề luyện thi {môn}'", "'NHIỆM VỤ NGÀY'"],
        ],
    )

    h(doc, "B2. Luồng Frontend (Student)", 2)
    p(doc, "B2.1. Luyện đề tự do", bold=True)
    num(doc, "Route: /student/quizzes — quizzes-client.tsx, generate-quiz-modal.tsx")
    num(doc, "Hook: useGenerateQuiz() — form subject_id, topic, difficulty, total_questions")
    num(doc, "Gọi quizService.generate() → POST /quizzes/generate (timeout 120s)")
    num(doc, "Thành công → router.push(/student/quizzes/{id})")
    num(doc, "Hook: useQuizAttempt(quizId) — load quiz, đếm thời gian, chọn đáp án, submit")
    num(doc, "Submit → POST /quizzes/{id}/submit → reload quiz ở chế độ review (có đáp án)")

    p(doc, "B2.2. Nhiệm vụ ngày (Quick Quiz)", bold=True)
    num(doc, "Component: quick-quiz-pane.tsx trong task-study-view / task-study-modal")
    num(doc, "Props: subjectId, topic, studyPlanId")
    num(doc, "Bước 1: GET /quizzes/plan/{studyPlanId} — nếu đã có đề thì load luôn")
    num(doc, "Bước 2 (404): POST /quizzes/generate kèm study_plan_id")
    num(doc, "Làm bài inline → submit → hiển thị result (score, correct/wrong)")
    num(doc, "onSuccess() → refresh study plan (đánh dấu hoàn thành nếu ≥8)")

    p(doc, "B2.3. Lịch sử & Review", bold=True)
    num(doc, "GET /quizzes/student/history → quiz-history-table.tsx")
    num(doc, "GET /quizzes/{id}/review — chỉ sau khi đã submit (permission check BE)")

    h(doc, "B3. API Reference — Quizzes", 2)
    tbl(
        doc,
        ["Method", "Endpoint", "Auth", "Handler / Service", "Response"],
        [
            ["POST", "/api/v1/quizzes/generate", "Student", "generate_and_save_quiz", "QuizResponse (ẩn đáp án)"],
            ["GET", "/api/v1/quizzes/{quiz_id}", "User", "get_quiz", "QuizResponse"],
            ["GET", "/api/v1/quizzes/{quiz_id}/review", "User", "get_quiz_review", "QuizDetailResponse (full đáp án)"],
            ["GET", "/api/v1/quizzes/plan/{study_plan_id}", "User", "get_quiz_for_study_plan", "QuizResponse"],
            ["POST", "/api/v1/quizzes/{quiz_id}/submit", "Student", "submit_student_quiz + BG analytics", "QuizAttemptResponse"],
            ["GET", "/api/v1/quizzes/student/history", "Student", "get_student_quiz_attempts", "List history"],
            ["POST", "/api/v1/quizzes/teacher/create", "Teacher", "teacher_create_quiz", "QuizResponse"],
            ["GET", "/api/v1/quizzes/classroom/{id}/attempts", "Teacher", "get_classroom_quiz_attempts", "List attempts"],
        ],
    )

    p(doc, "Request body POST /quizzes/generate:", bold=True)
    code(
        doc,
        '{\n  "subject_id": 1,\n  "topic": "Thì hiện tại đơn",\n'
        '  "difficulty": "medium",\n  "total_questions": 5,\n'
        '  "study_plan_id": null  // hoặc ID nhiệm vụ ngày\n}',
    )

    p(doc, "Request body POST /quizzes/{id}/submit:", bold=True)
    code(
        doc,
        '{\n  "answers": [\n    {"question_index": 0, "answer": "B"},\n'
        '    {"question_index": 1, "answer": "A"}\n  ],\n'
        '  "duration_seconds": 245\n}',
    )

    h(doc, "B4. Luồng Backend sinh đề — từng bước", 2)
    steps_gen = [
        "API quizzes.py → generate_new_quiz() nhận QuizCreateRequest",
        "Inject: db (MySQL Session), db_mongo (MongoDB), current_user",
        "generate_and_save_quiz() — generation.py",
        "subject_repository.get(subject_id) — 404 ValueError nếu không có môn",
        "vector_search_materials(db_mongo, query_text=topic, subject_id, top_k=3)",
        "build_rag_context(materials) — ghép chuỗi context cho LLM",
        "_generate_with_qc_review(subject_name, topic, difficulty, total_questions, context, skip_qc)",
        "normalize_ai_questions(ai_quiz) — map correct_answer → key A/B/C/D",
        "Xử lý title: fallback nếu rỗng hoặc 'QuizResponse'",
        "quiz_repository.stage_ai_generated() — INSERT quizzes",
        "commit_or_rollback(db) + db.refresh()",
        "Trả Quiz ORM → Pydantic QuizResponse (ẩn correct_answer ở schema student)",
    ]
    for i, s in enumerate(steps_gen, 1):
        num(doc, f"B4.{i}. {s}")

    h(doc, "B5. RAG — Vector Search tài liệu", 2)
    p(doc, "File: app/services/embedding_service.py → vector_search_materials()", bold=True)
    tbl(
        doc,
        ["Tham số", "Giá trị", "Ghi chú"],
        [
            ["query_text", "topic (string)", "Embed câu query — KHÔNG dùng tên môn"],
            ["subject_id", "int", "Lọc chunk cùng môn học"],
            ["top_k", "3", "Lấy 3 đoạn liên quan nhất"],
            ["min_score", "0.55", "Cosine similarity tối thiểu"],
        ],
    )
    p(doc, "Luồng nội bộ RAG:", bold=True)
    num(doc, "Redis cache key: rag:{subject_id}:{hash(topic)} — TTL 3600s")
    num(doc, "get_embedding(query_text) — thread pool, model embedding")
    num(doc, "MongoDB Atlas $vectorSearch (ưu tiên) hoặc NumPy cosine fallback")
    num(doc, "Trả List[{topic, content, score, document_id, ...}]")
    num(doc, "build_rag_context() format: '--- Tài liệu N (Chủ đề: X) ---\\n{content}'")

    h(doc, "B6. Multi-Agent QC (Agent-to-Agent)", 2)
    p(
        doc,
        "Đây là mô hình Agent-to-Agent dạng pipeline: Service orchestrate, "
        "không phải agent tự gọi agent. Chỉ 1 vòng review → correct, không lặp QC lần 2.",
    )
    tbl(
        doc,
        ["Bước", "Agent", "File", "Input", "Output"],
        [
            ["1", "Generator", "agents/quiz_generator/agent.py", "context + topic + subject", "QuizResponse JSON"],
            ["2", "QC Reviewer", "agents/quiz_generator/reviewer.py", "context + quiz_dict", "is_valid, feedback, error_indices"],
            ["3 (nếu fail)", "Generator correct", "agent.py correct_quiz_questions", "quiz + feedback + context", "QuizResponse đã sửa"],
        ],
    )
    p(doc, "4 tiêu chí QC Reviewer:", bold=True)
    bullet(doc, "Đáp án đúng khớp tài liệu gốc (context)")
    bullet(doc, "Không trùng lặp ý nghĩa giữa các câu")
    bullet(doc, "Giải thích (explanation) rõ ràng")
    bullet(doc, "Distractor (đáp án sai) hợp lý, không dễ loại trừ")

    p(doc, "Generator Agent — chi tiết generate_quiz():", bold=True)
    bullet(doc, "Prompt: QUIZ_GENERATOR_SYSTEM_PROMPT + RAG context + chống Prompt Injection")
    bullet(doc, "LLM: generate_content_nvidia, response_schema=QuizResponse, temperature 0.3→0.5 (retry)")
    bullet(doc, "Retry tối đa 3 lần; remove_duplicate_questions() bằng code")
    bullet(doc, "Chấp nhận nếu đủ total_questions hoặc ≥ min_acceptable = max(3, min(5, total))")

    p(doc, "Xử lý lỗi QC:", bold=True)
    bullet(doc, "QC crash (API lỗi) → fail-open: dùng đề bước 1, log warning")
    bullet(doc, "Parse JSON QC lỗi → is_valid=False → trigger correct")

    h(doc, "B7. Chuẩn hóa câu hỏi & Lưu DB", 2)
    p(doc, "normalize_ai_questions() — grading.py:", bold=True)
    bullet(doc, "Convert options → [{key, value}]")
    bullet(doc, "Nếu correct_answer không phải key hợp lệ → map theo value text hoặc fallback options[0]")
    bullet(doc, "Output lưu vào quizzes.questions JSON")

    h(doc, "B8. Luồng làm bài & Submit", 2)
    tbl(
        doc,
        ["Layer", "Hàm", "File"],
        [
            ["API", "submit_quiz()", "api/v1/quizzes.py"],
            ["Wrapper", "submit_student_quiz()", "services/quiz/queries.py"],
            ["Core", "submit_quiz_attempt()", "services/quiz/attempts.py"],
            ["Grading", "grade_submission()", "services/quiz/grading.py"],
            ["Repo", "attempt_repository.stage_attempt()", "repositories/attempt_repository.py"],
        ],
    )
    num(doc, "Kiểm tra role STUDENT — 403 nếu không phải học sinh")
    num(doc, "submit_quiz_attempt: load quiz, grade, stage attempt, commit")
    num(doc, "Nếu quiz.study_plan_id AND score >= 8.0 → plan_repository.mark_completed()")
    num(doc, "submit_student_quiz trả (attempt, subject_id) cho analytics")
    num(doc, "background_tasks.add_task(_analytics_background) — mở SessionLocal riêng")

    h(doc, "B9. Cơ chế chấm điểm", 2)
    p(doc, "grade_submission() — KHÔNG dùng AI, pure Python:", bold=True)
    num(doc, "submitted_map = {question_index: answer}")
    num(doc, "for idx, q in enumerate(questions): so sánh upper().strip()")
    num(doc, "Công thức: score = (correct_count / total) × 10")
    num(doc, "answers_json: [{question_index, answer, is_correct}, ...]")

    h(doc, "B10. Cấu trúc dữ liệu JSON", 2)
    p(doc, "quizzes.questions (lưu MySQL):", bold=True)
    code(
        doc,
        '[\n  {\n    "question_text": "Chọn thì đúng...",\n'
        '    "options": [{"key": "A", "value": "..."}, {"key": "B", "value": "..."}],\n'
        '    "correct_answer": "B",\n    "explanation": "Vì trong tài liệu..."\n  }\n]',
    )
    p(doc, "quiz_attempts.answers (sau chấm):", bold=True)
    code(
        doc,
        '[\n  {"question_index": 0, "answer": "A", "is_correct": false},\n'
        '  {"question_index": 1, "answer": "B", "is_correct": true}\n]',
    )

    h(doc, "B11. Bảng MySQL / MongoDB", 2)
    tbl(
        doc,
        ["Bảng / Collection", "Cột quan trọng", "Mô tả"],
        [
            ["quizzes", "id, student_id, subject_id, study_plan_id, questions JSON", "Đề thi"],
            ["quiz_attempts", "quiz_id, student_id, score, answers, duration_seconds", "Lần làm bài"],
            ["study_plans", "id, student_id, is_completed", "Nhiệm vụ ngày"],
            ["subjects", "id, name", "Môn học"],
            ["MongoDB embeddings", "subject_id, content, vector", "Chunk tài liệu cho RAG"],
        ],
    )

    h(doc, "B12. Xử lý lỗi & Timeout", 2)
    tbl(
        doc,
        ["Tình huống", "Xử lý"],
        [
            ["Generate LLM fail 3 lần", "RuntimeError 500"],
            ["Không tìm thấy môn", "400 ValueError"],
            ["RAG không có tài liệu", "context='' — AI sinh từ kiến thức chung"],
            ["QC crash", "Fail-open — dùng đề chưa QC"],
            ["Submit quiz không tồn tại", "404"],
            ["Xem review chưa làm bài", "403 PermissionError"],
            ["FE generate timeout", "120s — toast lỗi"],
        ],
    )

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN C — RECOMMENDATION
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN C — LUỒNG AI RECOMMENDATION", 1)

    h(doc, "C1. Điều kiện kích hoạt & Tách sync/async", 2)
    tbl(
        doc,
        ["Giai đoạn", "Sync/Async", "Thời điểm HS thấy"],
        [
            ["Chấm điểm + lưu attempt", "ĐỒNG BỘ", "Ngay — response API submit"],
            ["Mark study plan done (≥8)", "ĐỒNG BỘ", "Ngay sau submit (nhiệm vụ ngày)"],
            ["Recalculate analytics", "NGẦM", "Vài giây sau"],
            ["Analytics Agent", "NGẦM", "—"],
            ["Recommender (score<8)", "NGẦM", "Notification 'chờ duyệt'"],
            ["GV duyệt", "Thủ công", "HS thấy recommendations sau approve"],
        ],
    )
    p(doc, "Cơ chế ngầm: FastAPI BackgroundTasks (KHÔNG Celery). File: api/v1/quizzes.py", bold=True)

    h(doc, "C2. Background Task chi tiết", 2)
    p(doc, "update_student_analytics_and_recommend() — analytic_service.py", bold=True)
    steps_analytics = [
        "_load_submission_context(student_id, subject_id, quiz_id) — load subject, student, quiz+classroom",
        "_recalculate_learning_analytic() — đếm attempts, tính average_score, build attempts_history",
        "Upsert learning_analytics (quizzes_completed, average_score)",
        "_apply_ai_analytics() — asyncio.to_thread(evaluate_learning_performance)",
        "Ghi weak_topics, strong_topics, ai_feedback vào learning_analytics",
        "IF score < 8.0 → _create_pending_recommendation()",
        "IF score < 7.0 + có goal → _trigger_adaptive_plan()",
        "_add_progress_notification() — luôn gửi notification cập nhật hồ sơ",
        "commit_or_rollback(db)",
    ]
    for i, s in enumerate(steps_analytics, 1):
        num(doc, f"C2.{i}. {s}")

    h(doc, "C3. Analytics Agent", 2)
    p(doc, "File: app/agents/analytics/agent.py → evaluate_learning_performance()", bold=True)
    tbl(
        doc,
        ["Input", "Output field", "Mô tả"],
        [
            ["subject_name", "—", "Tên môn"],
            ["attempts_history[]", "weak_topics", "Chủ đề yếu — prompt: TB < 6.5"],
            ["{topic, score, is_passed}", "strong_topics", "Chủ đề mạnh — prompt: TB ≥ 8.0"],
            ["—", "learning_trend", "improving | declining | stable"],
            ["—", "ai_feedback", "Nhận xét tổng quan"],
        ],
    )
    p(doc, "is_passed trong history: score >= PASS_SCORE_THRESHOLD (5.0)", bold=True)

    h(doc, "C4. Recommender Agent", 2)
    p(doc, "Điều kiện: score < SCORE_RECOMMENDATION_THRESHOLD (8.0)", bold=True)
    p(doc, "File: app/agents/recommender/agent.py → generate_recommendation()", bold=True)
    bullet(doc, "Input: subject_name, topic_name (quiz.title), score, weak_topics từ Analytics")
    bullet(doc, "Output: văn bản markdown 100–200 từ — nhận xét, nội dung cần ôn, 2–3 hành động")
    bullet(doc, "Lưu: ai_recommendation_reviews — status=pending, teacher_id từ quiz.classroom")
    bullet(doc, "Notification HS: 'Đề xuất ôn tập AI đang chờ giáo viên duyệt'")

    h(doc, "C5. Human-in-the-Loop — Giáo viên duyệt", 2)
    p(doc, "Frontend: /teacher/recommendations — recommendation-review-card.tsx", bold=True)
    p(doc, "Backend: recommendation_service.py → review_recommendation()", bold=True)
    num(doc, "GET /recommendations/pending — danh sách pending theo lớp GV dạy")
    num(doc, "PATCH /recommendations/{review_id} — body: {status: approved|rejected, teacher_feedback?}")
    num(doc, "Kiểm tra quyền: teacher_teaches_student()")
    p(doc, "Khi APPROVED:", bold=True)
    bullet(doc, "stage_review_plan — ngày mai 19:00–20:00 (nếu HS có goal)")
    bullet(doc, "Notification HS: 'Đề xuất học tập mới được phê duyệt'")
    bullet(doc, "Invalidate Redis cache recommendations")
    bullet(doc, "Email SMTP (async) — send_recommendation_approved_email")
    p(doc, "Khi REJECTED:", bold=True)
    bullet(doc, "Không hiện trong /my-recommendations")
    bullet(doc, "Email từ chối kèm teacher_feedback (nếu có)")

    h(doc, "C6. Học sinh nhận kết quả", 2)
    tbl(
        doc,
        ["Thời điểm", "Kênh", "Nội dung"],
        [
            ["Ngay submit", "API response", "score, correct_count, wrong_count"],
            ["~vài giây", "notifications", "Hồ sơ học tập được cập nhật"],
            ["score<8", "notifications", "Đề xuất chờ GV duyệt"],
            ["Sau approve", "/student/recommendations", "GET /my-recommendations (approved only)"],
            ["Sau approve", "study_plans", "Plan ôn tập ngày mai (nếu có goal)"],
            ["Sau approve", "email", "Nội dung đề xuất"],
        ],
    )
    p(doc, "Cache HS recommendations: Redis TTL 600s, invalidate khi GV duyệt.", bold=True)

    h(doc, "C7. Adaptive Plan (score < 7)", 2)
    p(doc, "_trigger_adaptive_plan() — analytic_service.py", bold=True)
    bullet(doc, "Điều kiện: score < 7.0 AND goal active AND weak_topics không rỗng")
    bullet(doc, "Tạo chat session MongoDB: 'Tự động điều chỉnh lộ trình - {môn}'")
    bullet(doc, "add_chat_message(user) mô tả điểm yếu → Roadmap Planner Agent xử lý")
    bullet(doc, "Notification: 'Lộ trình học được điều chỉnh tự động'")

    h(doc, "C8. Notification, Email, Cache", 2)
    tbl(
        doc,
        ["Loại", "NotificationType", "Ai nhận"],
        [
            ["Cập nhật hồ sơ", "SCORE", "Học sinh"],
            ["Chờ duyệt đề xuất", "PLAN", "Học sinh"],
            ["Đề xuất approved", "PLAN", "Học sinh"],
            ["Adaptive plan", "PLAN", "Học sinh"],
        ],
    )

    h(doc, "C9. API Reference — Recommendations", 2)
    tbl(
        doc,
        ["Method", "Endpoint", "Role", "Service"],
        [
            ["GET", "/api/v1/recommendations/pending", "Teacher", "get_pending_reviews"],
            ["GET", "/api/v1/recommendations/my-recommendations", "Student", "get_student_recommendations"],
            ["PATCH", "/api/v1/recommendations/{review_id}", "Teacher", "review_recommendation"],
        ],
    )

    h(doc, "C10. Cấu trúc dữ liệu & Bảng DB", 2)
    p(doc, "ai_recommendation_reviews:", bold=True)
    tbl(
        doc,
        ["Cột", "Kiểu", "Mô tả"],
        [
            ["id", "BigInteger PK", "ID"],
            ["student_id", "FK users", "Học sinh"],
            ["teacher_id", "FK users nullable", "GV duyệt"],
            ["recommendation", "Text", "Nội dung AI"],
            ["teacher_feedback", "Text nullable", "Ghi chú GV"],
            ["status", "Enum", "pending | approved | rejected"],
            ["created_at", "DateTime", "Thời gian tạo"],
        ],
    )
    p(doc, "learning_analytics:", bold=True)
    tbl(
        doc,
        ["Cột", "Mô tả"],
        [
            ["student_id + subject_id", "Unique constraint"],
            ["average_score", "TB tất cả attempts theo môn"],
            ["quizzes_completed", "Số lần làm bài"],
            ["weak_topics / strong_topics", "JSON từ Analytics Agent"],
            ["ai_feedback", "Text nhận xét AI"],
        ],
    )

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN D — 50 BƯỚC END-TO-END
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN D — LUỒNG END-TO-END (50 BƯỚC)", 1)
    p(doc, "Kịch bản: Học sinh luyện đề tự do → làm bài → điểm thấp → nhận khuyến nghị sau GV duyệt.")

    e2e = [
        "[FE] HS mở /student/quizzes, bấm 'Sinh đề AI'",
        "[FE] useGenerateQuiz — nhập môn, topic, difficulty, số câu",
        "[FE] POST /quizzes/generate (timeout 120s)",
        "[BE] generate_new_quiz — auth current_user",
        "[BE] generate_and_save_quiz — load subject",
        "[BE] vector_search_materials(topic, subject_id) — Redis/Mongo RAG",
        "[BE] build_rag_context — format tài liệu",
        "[BE] quiz_generator.generate — Generator Agent (LLM call 1)",
        "[BE] remove_duplicate_questions — lọc trùng code",
        "[BE] quiz_generator.review — QC Reviewer Agent (LLM call 2)",
        "[BE] Nếu is_valid=False → quiz_generator.correct (LLM call 3)",
        "[BE] normalize_ai_questions — chuẩn hóa A/B/C/D",
        "[BE] INSERT quizzes — commit MySQL",
        "[FE] Redirect /student/quizzes/{id}",
        "[FE] useQuizAttempt — GET /quizzes/{id} (ẩn đáp án)",
        "[FE] HS chọn đáp án, timer duration tăng mỗi giây",
        "[FE] POST /quizzes/{id}/submit + answers + duration",
        "[BE] submit_quiz — kiểm tra role STUDENT",
        "[BE] submit_quiz_attempt — load quiz",
        "[BE] grade_submission — vòng lặp so sánh key đáp án",
        "[BE] INSERT quiz_attempts — score, answers JSON",
        "[BE] (Nếu study_plan + score≥8) mark_completed — bỏ qua với luyện đề tự do",
        "[BE] Trả QuizAttemptResponse — HS thấy điểm",
        "[BE] background_tasks.add_task(_analytics_background)",
        "[BE] SessionLocal mới → update_student_analytics_and_recommend",
        "[BE] _recalculate_learning_analytic — average_score, quizzes_completed",
        "[BE] Analytics Agent — weak/strong topics (LLM call 4)",
        "[BE] UPDATE learning_analytics",
        "[BE] score < 8? → Recommender Agent (LLM call 5)",
        "[BE] INSERT ai_recommendation_reviews status=pending",
        "[BE] Notification HS — chờ GV duyệt",
        "[BE] score < 7? → chat adaptive plan (nếu có goal)",
        "[BE] Notification HS — hồ sơ cập nhật",
        "[BE] commit analytics session",
        "[FE/GV] GV mở /teacher/recommendations",
        "[FE/GV] GET /recommendations/pending",
        "[FE/GV] Xem nội dung recommendation + thông tin HS",
        "[FE/GV] Bấm Approve + feedback (tùy chọn)",
        "[BE] PATCH /recommendations/{id} status=approved",
        "[BE] Kiểm tra teacher_teaches_student",
        "[BE] stage_review_plan — ngày mai 19h (nếu có goal)",
        "[BE] Notification HS — đã phê duyệt",
        "[BE] Invalidate Redis cache recommendations",
        "[BE] Email approved → HS",
        "[FE/HS] Mở /student/recommendations",
        "[FE/HS] GET /my-recommendations — chỉ approved",
        "[FE/HS] Hiển thị recommendation-card",
        "[FE/HS] (Tuỳ chọn) Xem study plan ôn tập ngày mai",
        "[FE/HS] GET /quizzes/student/history — xem điểm các lần làm",
        "[FE/HS] GET /quizzes/{id}/review — xem đáp án + giải thích",
        "[Kết thúc] Chu trình hoàn tất",
    ]
    for i, step in enumerate(e2e, 1):
        num(doc, f"D.{i}. {step}")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN E — FILE INDEX
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN E — CHỈ MỤC FILE MÃ NGUỒN", 1)
    tbl(
        doc,
        ["File", "Luồng", "Vai trò"],
        [
            ["Frontend/.../services/quiz.ts", "Quiz", "API client"],
            ["Frontend/.../hooks/use-generate-quiz.ts", "Quiz", "Sinh đề tự do"],
            ["Frontend/.../hooks/use-quiz-attempt.ts", "Quiz", "Làm bài + submit"],
            ["Frontend/.../quick-quiz-pane.tsx", "Quiz", "Nhiệm vụ ngày"],
            ["Frontend/.../services/recommendation.ts", "Rec", "API client HS"],
            ["Frontend/.../teacher/recommendations/page.tsx", "Rec", "GV duyệt"],
            ["Backend/app/api/v1/quizzes.py", "Quiz+Rec", "HTTP + BackgroundTasks"],
            ["Backend/app/api/v1/recommendations.py", "Rec", "HTTP recommendations"],
            ["Backend/app/services/quiz/generation.py", "Quiz", "Orchestrate sinh đề + QC"],
            ["Backend/app/services/quiz/grading.py", "Quiz", "Chấm + normalize"],
            ["Backend/app/services/quiz/attempts.py", "Quiz", "Submit + history"],
            ["Backend/app/services/quiz/agent_port.py", "Quiz", "Adapter agents"],
            ["Backend/app/agents/quiz_generator/agent.py", "Quiz", "Generator Agent"],
            ["Backend/app/agents/quiz_generator/reviewer.py", "Quiz", "QC Reviewer Agent"],
            ["Backend/app/services/embedding_service.py", "Quiz", "RAG vector search"],
            ["Backend/app/services/analytic_service.py", "Rec", "Post-submit analytics"],
            ["Backend/app/agents/analytics/agent.py", "Rec", "Analytics Agent"],
            ["Backend/app/agents/recommender/agent.py", "Rec", "Recommender Agent"],
            ["Backend/app/services/recommendation_service.py", "Rec", "HITL duyệt"],
        ],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN F — NGƯỠNG
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN F — BẢNG NGƯỠNG NGHIỆP VỤ", 1)
    tbl(
        doc,
        ["Ngưỡng", "Giá trị", "File", "Hành vi"],
        [
            ["Hoàn thành nhiệm vụ ngày", "≥ 8.0", "grading.py PLAN_PASS_SCORE_THRESHOLD", "mark_completed study_plan"],
            ["Kích hoạt recommendation", "< 8.0", "analytic_service SCORE_RECOMMENDATION_THRESHOLD", "Recommender + pending"],
            ["Adaptive plan", "< 7.0", "analytic_service SCORE_ADAPTIVE_PLAN_THRESHOLD", "Chat điều chỉnh lộ trình"],
            ["Pass trong history", "≥ 5.0", "PASS_SCORE_THRESHOLD", "is_passed=true trong analytics"],
            ["Chủ đề yếu (AI)", "< 6.5", "Analytics prompt", "weak_topics"],
            ["Chủ đề mạnh (AI)", "≥ 8.0", "Analytics prompt", "strong_topics"],
            ["RAG similarity", "≥ 0.55", "embedding_service min_score", "Lọc chunk"],
            ["RAG cache TTL", "3600s", "embedding_service", "Redis"],
            ["Rec cache TTL", "600s", "recommendation_service", "Redis per student"],
            ["FE generate timeout", "120s", "quiz.ts", "Axios"],
        ],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # PHẦN G — KẾT LUẬN
    # ═══════════════════════════════════════════════════════════════════════
    h(doc, "PHẦN G — KẾT LUẬN", 1)
    p(
        doc,
        "Luồng Luyện thi được thiết kế theo kiến trúc phân tầng rõ ràng: Frontend gọi REST API, "
        "Service orchestrate RAG và Multi-Agent QC (Agent-to-Agent pipeline), chấm điểm đồng bộ "
        "bằng vòng lặp so sánh deterministic. Luồng AI Recommendation tách biệt, chạy ngầm qua "
        "BackgroundTasks, kết hợp Analytics Agent và Recommender Agent với mô hình HITL — "
        "đảm bảo giáo viên kiểm soát chất lượng nội dung khuyến nghị trước khi học sinh nhận. "
        "Hai luồng liên kết tại điểm POST /quizzes/{id}/submit: phần đồng bộ trả điểm ngay, "
        "phần ngầm kích hoạt toàn bộ pipeline phân tích và đề xuất học tập.",
    )

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(para.add_run("— Hết báo cáo chi tiết —"), size=11, color=RGBColor(80, 80, 80))

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
