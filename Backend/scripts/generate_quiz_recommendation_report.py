"""
Generate Word report: Quiz practice flow + AI Recommendation flow.
Run: python Backend/scripts/generate_quiz_recommendation_report.py
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = (
    Path(__file__).resolve().parent.parent
    / "materials"
    / "Bao_cao_Luyen_de_va_Khuyen_nghi_AI.docx"
)


def set_run_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, bold=True, size=16)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    doc.add_paragraph()


def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    add_title(doc, "BÁO CÁO LUỒNG NGHIỆP VỤ")
    add_title(doc, "Luyện đề trắc nghiệm AI & Khuyến nghị học tập AI")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Nền tảng AI Learning Assistant Platform\n"
        f"Ngày lập: {datetime.now().strftime('%d/%m/%Y')}"
    )
    set_run_font(run, size=12)
    doc.add_page_break()

    # ── MỤC LỤC NỘI DUNG (manual) ───────────────────────────────────────────
    add_heading(doc, "MỤC LỤC", 1)
    for item in [
        "1. Tổng quan hệ thống",
        "2. Phần I — Luồng luyện đề trắc nghiệm AI",
        "   2.1. Mục tiêu nghiệp vụ",
        "   2.2. Kiến trúc tầng xử lý",
        "   2.3. Luồng sinh đề (Generate)",
        "   2.4. Cơ chế RAG — tìm tài liệu",
        "   2.5. Multi-Agent QC Review",
        "   2.6. Luồng làm bài và nộp bài (Submit)",
        "   2.7. Cơ chế chấm điểm tự động",
        "   2.8. Bảng dữ liệu & API liên quan",
        "3. Phần II — Luồng khuyến nghị AI (HITL)",
        "   3.1. Mục tiêu nghiệp vụ",
        "   3.2. Kích hoạt sau submit (Background Tasks)",
        "   3.3. Analytics Agent — phân tích học lực",
        "   3.4. Recommender Agent — sinh đề xuất ôn tập",
        "   3.5. Human-in-the-Loop — giáo viên duyệt",
        "   3.6. Học sinh nhận kết quả",
        "   3.7. Luồng điều chỉnh lộ trình (Adaptive Plan)",
        "   3.8. Bảng dữ liệu & API liên quan",
        "4. Ngưỡng nghiệp vụ (Business Thresholds)",
        "5. Sơ đồ tổng hợp",
        "6. Kết luận",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # ── 1. TỔNG QUAN ───────────────────────────────────────────────────────
    add_heading(doc, "1. Tổng quan hệ thống", 1)
    add_para(
        doc,
        "Hệ thống AI Learning Assistant Platform hỗ trợ học sinh luyện tập trắc nghiệm "
        "theo tài liệu giáo trình (RAG) và tự động đề xuất nội dung ôn tập khi kết quả "
        "chưa đạt yêu cầu. Hai luồng nghiệp vụ chính được mô tả trong báo cáo này:",
    )
    add_bullet(doc, "Luồng I — Luyện đề: sinh đề AI → học sinh làm bài → chấm điểm tự động.")
    add_bullet(
        doc,
        "Luồng II — Khuyến nghị AI: phân tích kết quả ngầm → sinh đề xuất → "
        "giáo viên duyệt (HITL) → học sinh nhận đề xuất chính thức.",
    )
    add_para(
        doc,
        "Công nghệ backend: FastAPI, MySQL (dữ liệu chính), MongoDB (embedding tài liệu & chat), "
        "Redis (cache), NVIDIA NIM (LLM). Frontend: Next.js.",
    )

    # ── PHẦN I ─────────────────────────────────────────────────────────────
    add_heading(doc, "2. Phần I — Luồng luyện đề trắc nghiệm AI", 1)

    add_heading(doc, "2.1. Mục tiêu nghiệp vụ", 2)
    add_para(
        doc,
        "Cho phép học sinh tạo đề luyện tập trắc nghiệm (MCQ) theo môn học và chủ đề, "
        "bám sát tài liệu đã upload. Hệ thống sinh câu hỏi bằng AI, kiểm tra chất lượng (QC), "
        "lưu đề vào cơ sở dữ liệu, cho phép học sinh làm bài và chấm điểm tự động thang 0–10.",
    )

    add_heading(doc, "2.2. Kiến trúc tầng xử lý", 2)
    add_table(
        doc,
        ["Tầng", "Thành phần", "Vai trò"],
        [
            ["API", "app/api/v1/quizzes.py", "HTTP endpoints, auth, BackgroundTasks"],
            ["Service", "app/services/quiz/", "Nghiệp vụ quiz (tách theo use case)"],
            ["Agent", "app/agents/quiz_generator/", "Sinh đề, QC Reviewer, sửa đề"],
            ["Infrastructure", "embedding_service, nim_client", "RAG vector search, gọi LLM"],
            ["Repository", "quiz_repository, attempt_repository", "Truy cập MySQL"],
        ],
    )
    add_para(doc, "Cấu trúc package app/services/quiz/:", bold=True)
    add_table(
        doc,
        ["File", "Trách nhiệm"],
        [
            ["generation.py", "Orchestrate RAG + sinh đề + QC + lưu DB"],
            ["agent_port.py", "Adapter tới Quiz Generator Agent (dễ mock test)"],
            ["grading.py", "Chuẩn hóa câu hỏi AI + chấm điểm vòng lặp"],
            ["attempts.py", "Nộp bài, lưu attempt, hoàn thành nhiệm vụ ngày"],
            ["queries.py", "Đọc đề, quyền xem đáp án, wrapper submit"],
            ["teacher.py", "Giáo viên tạo đề thủ công (không qua AI)"],
        ],
    )

    add_heading(doc, "2.3. Luồng sinh đề (Generate)", 2)
    add_para(doc, "Endpoint: POST /api/v1/quizzes/generate", bold=True)
    add_numbered(doc, "Frontend gửi: subject_id, topic, difficulty, total_questions, study_plan_id (tùy chọn).")
    add_numbered(doc, "API gọi generate_and_save_quiz() trong generation.py.")
    add_numbered(doc, "Kiểm tra môn học tồn tại (subject_repository).")
    add_numbered(doc, "vector_search_materials() — tìm top 3 tài liệu liên quan trong MongoDB.")
    add_numbered(doc, "build_rag_context() — ghép nội dung tài liệu thành chuỗi context.")
    add_numbered(doc, "_generate_with_qc_review() — sinh đề + QC (xem mục 2.5).")
    add_numbered(doc, "normalize_ai_questions() — chuẩn hóa correct_answer thành key A/B/C/D.")
    add_numbered(doc, "quiz_repository.stage_ai_generated() — lưu bảng quizzes, commit MySQL.")
    add_numbered(doc, "Trả về QuizResponse cho frontend (ẩn đáp án đúng).")

    add_heading(doc, "2.4. Cơ chế RAG — tìm tài liệu", 2)
    add_para(doc, "Hàm: app/services/embedding_service.py → vector_search_materials()", bold=True)
    add_table(
        doc,
        ["Tham số", "Giá trị", "Mô tả"],
        [
            ["query_text", "topic", "Chủ đề HS nhập — dùng embed tìm kiếm"],
            ["subject_id", "ID môn học", "Lọc tài liệu cùng môn"],
            ["top_k", "3", "Lấy 3 chunk liên quan nhất"],
            ["min_score", "0.55", "Ngưỡng cosine similarity tối thiểu"],
        ],
    )
    add_para(
        doc,
        "Luồng search: kiểm tra Redis cache → embed query → MongoDB Atlas $vectorSearch "
        "(fallback NumPy cosine) → cache kết quả 1 giờ. "
        "Tên môn học (subject.name) được dùng khi gọi AI sinh câu hỏi, không dùng trực tiếp trong vector search.",
    )

    add_heading(doc, "2.5. Multi-Agent QC Review", 2)
    add_para(
        doc,
        "QC chỉ chạy với luyện đề tự do (study_plan_id = null). "
        "Nhiệm vụ ngày gắn study_plan_id → skip_qc=True để sinh nhanh hơn.",
    )
    add_para(doc, "Ba bước trong _generate_with_qc_review():", bold=True)
    add_numbered(
        doc,
        "Generator Agent (generate_quiz): LLM sinh bộ câu hỏi MCQ từ RAG context. "
        "Retry tối đa 3 lần, lọc trùng câu hỏi bằng code (remove_duplicate_questions).",
    )
    add_numbered(
        doc,
        "QC Reviewer Agent (review_generated_quiz): LLM thứ hai đọc context + đề, "
        "đánh giá 4 tiêu chí: (1) đáp án đúng khớp tài liệu, (2) không trùng ý, "
        "(3) giải thích rõ, (4) distractor hợp lý. Trả về is_valid, feedback, error_question_indices.",
    )
    add_numbered(
        doc,
        "Nếu is_valid=False → correct_quiz_questions(): Generator sửa lại toàn bộ đề theo feedback QC. "
        "Chỉ sửa một vòng, không gọi Reviewer lần 2.",
    )
    add_para(
        doc,
        "Xử lý lỗi: nếu QC crash (API lỗi) → bỏ qua QC, vẫn dùng đề bước 1 (fail-open). "
        "Nếu parse JSON QC lỗi → coi is_valid=False và trigger sửa.",
    )

    add_heading(doc, "2.6. Luồng làm bài và nộp bài (Submit)", 2)
    add_table(
        doc,
        ["Bước", "Endpoint / Hàm", "Mô tả"],
        [
            ["Lấy đề", "GET /quizzes/{id}", "get_quiz() — ẩn đáp án đúng"],
            ["Lấy đề theo lịch", "GET /quizzes/plan/{plan_id}", "get_quiz_for_study_plan()"],
            ["Nộp bài", "POST /quizzes/{id}/submit", "submit_student_quiz() → submit_quiz_attempt()"],
            ["Xem đáp án", "GET /quizzes/{id}/review", "Sau khi đã làm bài"],
            ["Lịch sử", "GET /quizzes/student/history", "Danh sách attempt + điểm"],
        ],
    )
    add_para(doc, "Payload submit: answers [{question_index, answer}], duration_seconds.", bold=True)

    add_heading(doc, "2.7. Cơ chế chấm điểm tự động", 2)
    add_para(doc, "Hàm: grade_submission() trong app/services/quiz/grading.py", bold=True)
    add_numbered(doc, "Map câu trả lời HS: {question_index → answer}.")
    add_numbered(
        doc,
        "Vòng lặp từng câu: so sánh str(chosen).strip().upper() == str(correct).strip().upper().",
    )
    add_numbered(doc, "Câu trống → tính sai. Không dùng AI chấm lại, không fuzzy match.")
    add_numbered(doc, "score = (correct_count / total_questions) × 10 — thang 0–10.")
    add_numbered(doc, "Lưu quiz_attempts: score, correct_count, wrong_count, answers JSON.")
    add_numbered(
        doc,
        "Nếu quiz gắn study_plan_id và score ≥ 8.0 → mark_completed() nhiệm vụ ngày (đồng bộ).",
    )
    add_para(
        doc,
        "Sau khi trả response cho HS, API đăng ký Background Task analytics "
        "(xem Phần II) — không chặn response.",
    )

    add_heading(doc, "2.8. Bảng dữ liệu & API liên quan", 2)
    add_table(
        doc,
        ["Bảng MySQL", "Nội dung"],
        [
            ["quizzes", "Đề thi: questions JSON, title, difficulty, subject_id, study_plan_id"],
            ["quiz_attempts", "Lần làm bài: score, answers, duration, submitted_at"],
            ["study_plans", "Nhiệm vụ ngày — mark done khi score ≥ 8"],
            ["subjects", "Môn học"],
        ],
    )
    add_table(
        doc,
        ["MongoDB", "Nội dung"],
        [
            ["study_materials / embeddings", "Chunk tài liệu + vector cho RAG"],
        ],
    )

    doc.add_page_break()

    # ── PHẦN II ────────────────────────────────────────────────────────────
    add_heading(doc, "3. Phần II — Luồng khuyến nghị AI (Human-in-the-Loop)", 1)

    add_heading(doc, "3.1. Mục tiêu nghiệp vụ", 2)
    add_para(
        doc,
        "Sau khi học sinh nộp bài, hệ thống tự động cập nhật hồ sơ học tập, phân tích điểm mạnh/yếu "
        "bằng AI. Nếu điểm chưa đạt ngưỡng, sinh đề xuất ôn tập cụ thể và gửi cho giáo viên duyệt "
        "trước khi học sinh nhận đề xuất chính thức — mô hình Human-in-the-Loop (HITL).",
    )

    add_heading(doc, "3.2. Kích hoạt sau submit (Background Tasks)", 2)
    add_para(doc, "File: app/api/v1/quizzes.py → submit_quiz()", bold=True)
    add_numbered(doc, "Đồng bộ: submit_student_quiz() — chấm + lưu attempt (+ hoàn thành plan nếu đủ điều kiện).")
    add_numbered(
        doc,
        "Ngầm: background_tasks.add_task(_analytics_background) — mở Session DB riêng, "
        "gọi update_student_analytics_and_recommend() trong analytic_service.py.",
    )
    add_para(
        doc,
        "Lưu ý: dùng FastAPI BackgroundTasks (cùng process uvicorn), không phải Celery queue. "
        "HS nhận điểm ngay; analytics chạy sau vài giây.",
    )

    add_heading(doc, "3.3. Analytics Agent — phân tích học lực", 2)
    add_para(doc, "File: app/services/analytic_service.py", bold=True)
    add_para(doc, "Luồng update_student_analytics_and_recommend() — luôn chạy:", bold=True)
    add_numbered(
        doc,
        "_recalculate_learning_analytic(): đếm quizzes_completed, tính average_score "
        "từ tất cả quiz_attempts theo môn → upsert bảng learning_analytics.",
    )
    add_numbered(
        doc,
        "_apply_ai_analytics(): gọi evaluate_learning_performance() (Analytics Agent) "
        "với lịch sử làm bài → ghi weak_topics, strong_topics, ai_feedback.",
    )
    add_para(doc, "Analytics Agent (app/agents/analytics/agent.py):", bold=True)
    add_bullet(doc, "Input: tên môn + attempts_history [{topic, score, is_passed}].")
    add_bullet(doc, "Output JSON: weak_topics, strong_topics, learning_trend, ai_feedback.")
    add_bullet(doc, "Ngưỡng trong prompt: yếu < 6.5, mạnh ≥ 8.0.")

    add_heading(doc, "3.4. Recommender Agent — sinh đề xuất ôn tập", 2)
    add_para(doc, "Điều kiện: score < 8.0 (SCORE_RECOMMENDATION_THRESHOLD)", bold=True)
    add_para(doc, "Hàm: _create_pending_recommendation() → generate_recommendation()", bold=True)
    add_numbered(
        doc,
        "Recommender Agent (app/agents/recommender/agent.py) sinh văn bản đề xuất 100–200 từ: "
        "nhận xét kết quả, nội dung cần ôn, 2–3 hành động học tập cụ thể.",
    )
    add_numbered(
        doc,
        "Input: subject_name, topic_name (quiz.title), score, weak_topics từ Analytics.",
    )
    add_numbered(
        doc,
        "Lưu ai_recommendation_reviews với status=pending, teacher_id từ quiz.classroom.",
    )
    add_numbered(
        doc,
        "Gửi notification cho HS: 'Đề xuất ôn tập AI đang chờ giáo viên duyệt'.",
    )

    add_heading(doc, "3.5. Human-in-the-Loop — giáo viên duyệt", 2)
    add_table(
        doc,
        ["API", "Vai trò"],
        [
            ["GET /recommendations/pending", "GV xem danh sách chờ duyệt"],
            ["PATCH /recommendations/{id}", "GV approve hoặc reject + feedback"],
        ],
    )
    add_para(doc, "File: app/services/recommendation_service.py → review_recommendation()", bold=True)
    add_para(doc, "Khi APPROVED:", bold=True)
    add_bullet(doc, "Cập nhật status=approved, lưu teacher_feedback.")
    add_bullet(
        doc,
        "Nếu HS có goal active → tạo study_plan ôn tập ngày mai 19:00–20:00 "
        "(plan_repository.stage_review_plan).",
    )
    add_bullet(doc, "Notification cho HS: 'Đề xuất học tập mới được phê duyệt'.")
    add_bullet(doc, "Gửi email (SMTP) nếu cấu hình.")
    add_bullet(doc, "Invalidate Redis cache recommendations của HS.")
    add_para(doc, "Khi REJECTED:", bold=True)
    add_bullet(doc, "HS không thấy trong danh sách recommendations.")
    add_bullet(doc, "Có thể gửi email từ chối kèm feedback GV.")

    add_heading(doc, "3.6. Học sinh nhận kết quả", 2)
    add_table(
        doc,
        ["Thời điểm", "HS thấy gì", "Nơi xem"],
        [
            ["Ngay sau submit", "Điểm bài thi", "UI làm bài / toast"],
            ["Sau analytics (~vài giây)", "Notification cập nhật hồ sơ", "Dashboard / SSE"],
            ["Sau analytics, score<8", "Notification chờ GV duyệt", "Notifications"],
            ["Sau GV approve", "Đề xuất chính thức", "GET /recommendations/my-recommendations → /student/recommendations"],
            ["Sau GV approve", "Plan ôn tập (nếu có goal)", "Lịch học / study plans"],
        ],
    )
    add_para(
        doc,
        "Quan trọng: trang /student/recommendations chỉ hiển thị bản ghi status=approved. "
        "Pending không hiện ở đây — HS chỉ biết qua notification.",
    )

    add_heading(doc, "3.7. Luồng điều chỉnh lộ trình (Adaptive Plan)", 2)
    add_para(doc, "Điều kiện: score < 7.0 AND có goal active AND có weak_topics", bold=True)
    add_para(doc, "Hàm: _trigger_adaptive_plan() trong analytic_service.py", bold=True)
    add_bullet(doc, "Tạo chat session MongoDB: 'Tự động điều chỉnh lộ trình - {môn}'.")
    add_bullet(doc, "Thêm message user mô tả điểm yếu → Roadmap Planner Agent xử lý trong chat.")
    add_bullet(doc, "Notification: 'Lộ trình học được điều chỉnh tự động'.")

    add_heading(doc, "3.8. Bảng dữ liệu & API liên quan", 2)
    add_table(
        doc,
        ["Bảng MySQL", "Nội dung"],
        [
            ["learning_analytics", "average_score, quizzes_completed, weak/strong_topics, ai_feedback"],
            ["ai_recommendation_reviews", "recommendation text, status, teacher_id, student_id"],
            ["notifications", "Thông báo HS/GV"],
            ["study_plans", "Plan ôn tập sau khi GV approve"],
        ],
    )

    doc.add_page_break()

    # ── NGƯỠNG ─────────────────────────────────────────────────────────────
    add_heading(doc, "4. Ngưỡng nghiệp vụ (Business Thresholds)", 1)
    add_table(
        doc,
        ["Ngưỡng", "Giá trị", "File / Hằng số", "Ý nghĩa"],
        [
            ["Đạt nhiệm vụ ngày", "≥ 8.0", "PLAN_PASS_SCORE_THRESHOLD (grading.py)", "Mark study plan done"],
            ["Kích hoạt recommendation", "< 8.0", "SCORE_RECOMMENDATION_THRESHOLD (analytic_service.py)", "Sinh đề xuất pending"],
            ["Adaptive plan", "< 7.0", "SCORE_ADAPTIVE_PLAN_THRESHOLD", "Điều chỉnh lộ trình qua chat"],
            ["Pass trong analytics history", "≥ 5.0", "PASS_SCORE_THRESHOLD", "is_passed trong lịch sử AI"],
            ["Chủ đề yếu (AI prompt)", "< 6.5", "Analytics Agent prompt", "Phân loại weak_topics"],
            ["Chủ đề mạnh (AI prompt)", "≥ 8.0", "Analytics Agent prompt", "Phân loại strong_topics"],
            ["RAG similarity", "≥ 0.55", "vector_search_materials min_score", "Lọc chunk tài liệu"],
        ],
    )

    # ── SƠ ĐỒ ──────────────────────────────────────────────────────────────
    add_heading(doc, "5. Sơ đồ tổng hợp", 1)
    add_para(doc, "5.1. Luồng luyện đề", bold=True)
    add_para(
        doc,
        "POST /quizzes/generate\n"
        "  → vector_search(topic, subject_id)\n"
        "  → build_rag_context\n"
        "  → generate_quiz → [QC review → correct?]\n"
        "  → normalize_ai_questions → lưu quizzes\n"
        "  → HS làm bài → POST /quizzes/{id}/submit\n"
        "  → grade_submission (vòng lặp so sánh A/B/C/D)\n"
        "  → lưu quiz_attempts → (≥8 + plan?) hoàn thành nhiệm vụ\n"
        "  → Response điểm cho HS",
    )
    add_para(doc, "5.2. Luồng khuyến nghị AI", bold=True)
    add_para(
        doc,
        "BackgroundTasks (sau submit)\n"
        "  → recalculate learning_analytics\n"
        "  → Analytics Agent (weak/strong topics)\n"
        "  → [score < 8] Recommender Agent → ai_recommendation_reviews (pending)\n"
        "  → notification HS 'chờ duyệt'\n"
        "  → GV: GET /recommendations/pending\n"
        "  → GV: PATCH approve/reject\n"
        "  → [approve] notification + email + study_plan ngày mai\n"
        "  → HS: GET /recommendations/my-recommendations",
    )

    # ── KẾT LUẬN ───────────────────────────────────────────────────────────
    add_heading(doc, "6. Kết luận", 1)
    add_para(
        doc,
        "Luồng luyện đề và khuyến nghị AI được thiết kế tách biệt rõ ràng: "
        "phần chấm điểm và trả kết quả cho học sinh chạy đồng bộ, đảm bảo trải nghiệm tức thì; "
        "phần phân tích AI và đề xuất ôn tập chạy ngầm qua BackgroundTasks, "
        "kết hợp mô hình HITL để giáo viên kiểm soát chất lượng nội dung khuyến nghị trước khi "
        "học sinh nhận. Kiến trúc multi-agent (Generator + QC Reviewer cho sinh đề; "
        "Analytics + Recommender cho hậu kiểm) giúp nâng cao chất lượng đề thi và đề xuất học tập "
        "dựa trên tài liệu RAG và lịch sử làm bài thực tế.",
    )

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("— Hết báo cáo —")
    set_run_font(run, size=11, color=RGBColor(100, 100, 100))

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
