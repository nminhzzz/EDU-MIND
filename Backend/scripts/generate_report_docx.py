"""Convert the project's Markdown technical report into a styled Word document."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_inline(paragraph, text: str) -> None:
    """Render a compact subset of Markdown inline formatting."""
    token_pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for level, size in ((1, 18), (2, 15), (3, 12), (4, 11)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer
    footer_title = footer.paragraphs[0]
    footer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_title.add_run("AI Learning Assistant Platform — Báo cáo phân tích hệ thống")

    field_paragraph = footer.add_paragraph()
    field_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_paragraph.add_run("Trang ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_paragraph._p.append(field)


def add_cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO PHÂN TÍCH HỆ THỐNG")
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string(BLUE)

    product = document.add_paragraph()
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product_run = product.add_run("AI LEARNING ASSISTANT PLATFORM")
    product_run.bold = True
    product_run.font.size = Pt(20)

    document.add_paragraph()
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Công nghệ • Nghiệp vụ • Luồng người dùng • Cơ sở dữ liệu • Production")

    for _ in range(6):
        document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Phiên bản 1.0\nNgày rà soát mã nguồn: 16/08/2026").italic = True
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("MỤC LỤC", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Nhấn Ctrl+A rồi F9 trong Microsoft Word để cập nhật mục lục."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)
    document.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index in range(columns):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, value)
            if row_index == 0:
                shade_cell(cell, BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_BLUE)
    document.add_paragraph()


def add_code_block(document: Document, code: list[str], language: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 140, 120, 140)
    paragraph = cell.paragraphs[0]
    if language == "mermaid":
        label = paragraph.add_run("SƠ ĐỒ LUỒNG (MERMAID)\n")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string(BLUE)
    run = paragraph.add_run("\n".join(code))
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    document.add_paragraph()


def convert(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    add_cover(document)
    add_toc(document)

    index = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    skipped_source_title = False

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                add_code_block(document, code_lines, code_language)
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not skipped_source_title:
                skipped_source_title = True
            else:
                document.add_heading(text, level=min(level, 4))
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            run = paragraph.add_run(stripped.lstrip("> "))
            run.italic = True
            run.font.color.rgb = RGBColor(90, 90, 90)
            index += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, ordered.group(2))
        elif bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, stripped)
        index += 1

    core = document.core_properties
    core.title = "Báo cáo phân tích hệ thống AI Learning Assistant Platform"
    core.subject = "Công nghệ, nghiệp vụ, flow vai trò, cơ sở dữ liệu và production"
    core.author = "AI Learning Assistant Platform"
    core.keywords = "AI, Learning, FastAPI, Next.js, Celery, MySQL, MongoDB, Redis, RAG"
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate_report_docx.py INPUT.md OUTPUT.docx")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
